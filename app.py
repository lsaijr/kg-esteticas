import os, json, uuid, threading, calendar, secrets, hashlib
import cloudinary
import cloudinary.uploader
WEASYPRINT_OK = False
WeasyprintHTML = None
from datetime import date, datetime
from flask import Flask, request, jsonify, render_template_string, send_from_directory
import requests as req
import zipfile, re, html as htmllib

app = Flask(__name__)

# ── Config desde variables de entorno ────────────────────────
CLAUDE_KEY  = os.environ.get('CLAUDE_KEY', '')
GEMINI_KEY  = os.environ.get('GEMINI_KEY', '')
GROQ_KEY         = os.environ.get('GROQ_KEY', '')
OPENROUTER_KEY   = os.environ.get('OPENROUTER_KEY', '')
RESEND_KEY  = os.environ.get('RESEND_KEY', '')
MAIL_CC         = [m.strip() for m in os.environ.get('MAIL_CC','').split(',') if m.strip()]
ADMIN_PASSWORD  = os.environ.get('ADMIN_PASSWORD', 'carvajal2026')
MAIL_TO     = os.environ.get('MAIL_TO', 'isai.josue@gmail.com').strip()
MAIL_FROM   = os.environ.get('MAIL_FROM', 'envios@centrocarvajal.com')
DEMO_MAIL   = os.environ.get('DEMO_MAIL', 'isai.josue@gmail.com')
PLANES_DIR  = os.path.join(os.path.dirname(__file__), 'planes_generados')
os.makedirs(PLANES_DIR, exist_ok=True)

# ── Cloudinary ────────────────────────────────────────────────
CLOUDINARY_CLOUD_NAME = os.environ.get('CLOUDINARY_CLOUD_NAME', '')
CLOUDINARY_API_KEY    = os.environ.get('CLOUDINARY_API_KEY', '')
CLOUDINARY_API_SECRET = os.environ.get('CLOUDINARY_API_SECRET', '')

if CLOUDINARY_CLOUD_NAME:
    cloudinary.config(
        cloud_name = CLOUDINARY_CLOUD_NAME,
        api_key    = CLOUDINARY_API_KEY,
        api_secret = CLOUDINARY_API_SECRET,
        secure     = True
    )

# ── Jobs en memoria ───────────────────────────────────────────
jobs = {}  # jobId -> {'status': ..., 'msg': ..., 'html_url': ...}
admin_tokens = set()  # tokens de sesión activos

# ════════════════════════════════════════════════════════════
# RUTAS
# ════════════════════════════════════════════════════════════

@app.route('/', methods=['GET'])
def index():
    with open(os.path.join(os.path.dirname(__file__), 'index.html'), encoding='utf-8') as f:
        return f.read()

# ── Demo — formulario estético simplificado ──────────────────
@app.route('/demo')
def demo():
    with open(os.path.join(os.path.dirname(__file__), 'formulario-estetica-v2.html'), encoding='utf-8') as f:
        return f.read()

@app.route('/demo/recomendar', methods=['POST'])
def demo_recomendar():
    """Proxy multi-modelo para el formulario demo — las keys nunca salen del servidor."""
    try:
        data   = request.get_json(force=True)
        perfil = data.get('perfil', '')
        modelo = data.get('modelo', 'claude')
        if not perfil:
            return jsonify({'error': 'Sin datos de perfil'}), 400

        base_prompt = (
            'Eres una especialista en estética con años de experiencia. Tu forma de comunicarte es cálida, respetuosa y genuinamente personalizada. '
            'Le hablas al cliente por su nombre y usas "tú". '
            'TONO: profesional pero humano — como una especialista que realmente leyó su caso y le habla con consideración, no como un robot generando texto genérico ni como una amiga informal. '
            'EVITA absolutamente: frases genéricas como "¡Hola!", exclamaciones excesivas, emojis en el texto, frases vacías como "es un placer atenderte", lenguaje clínico frío, o recomendaciones que podrían ser para cualquier persona. '
            'BUSCA: oraciones que demuestren que leíste el perfil específico del cliente — menciona su situación real, sus áreas de interés, su nivel de estrés o actividad física cuando sea relevante. Que cada frase aporte algo concreto. '
            'Responde SIEMPRE en HTML simple usando solo estas etiquetas: <p>, <strong>, <ul>, <li>. '
            'NO uses markdown, NO uses asteriscos (**), NO uses encabezados, NO uses tablas. '
            'Para resaltar el nombre de un tratamiento usa ÚNICAMENTE la etiqueta HTML <strong>, nunca asteriscos. '
            'Estructura: '
            '1. Abre mencionando el nombre del cliente y una observación específica sobre su perfil que demuestre que lo leíste — algo como reconocer su preocupación principal o su situación particular. '
            '2. Explica brevemente por qué los tratamientos que vas a recomendar tienen sentido para SU caso específico, en lenguaje claro sin tecnicismos. '
            '3. Lista los tratamientos recomendados — nombre del tratamiento en negrita y una frase concreta sobre qué resultado puede esperar esta persona en particular. '
            '4. Cierra con una invitación a agendar su consulta, transmitiendo que en ese espacio podrán profundizar y resolver todas sus dudas.'
        )
        if modelo in ('groq', 'openrouter'):
            sys_prompt = base_prompt + (
                ' IMPORTANTE: sé detallada y muy personalizada, pero mantén un tono cercano y natural — nada de lenguaje corporativo ni frases de manual. '
                'Integra datos reales del perfil de forma orgánica en el texto: edad, nivel de estrés, actividad física, historial de tratamientos. '
                'No menciones el presupuesto bajo ningún concepto — ni directa ni indirectamente. Nunca uses frases como "considerando tu presupuesto", "dentro de tus posibilidades", "limitaciones económicas" o similares. '
                'Evita completamente frases hechas como "¿te gustaría explorar estas opciones?", "no dudes en contactarnos", "estamos a tu disposición", "el mejor curso de acción". '
                'Cada tratamiento merece 2-3 oraciones — qué hace, por qué encaja específicamente con este perfil y qué resultado concreto puede esperar esta persona. '
                'Cierra con una sola oración natural invitando a agendar, sin preguntas retóricas. '
                'Máximo 450 palabras.'
            )
        else:
            sys_prompt = base_prompt + ' Máximo 300 palabras. Cada frase debe aportar valor concreto — nada de relleno.'
        user_msg = 'Perfil del paciente:\n' + perfil

        # ── Claude ──────────────────────────────────────────
        if modelo == 'claude':
            r = req.post('https://api.anthropic.com/v1/messages',
                headers={'Content-Type':'application/json','x-api-key':CLAUDE_KEY,'anthropic-version':'2023-06-01'},
                json={'model':'claude-haiku-4-5-20251001','max_tokens':800,
                      'system':sys_prompt,'messages':[{'role':'user','content':user_msg}]},
                timeout=30)
            j = r.json()
            if 'content' in j and j['content']:
                return jsonify({'html': j['content'][0]['text']})
            return jsonify({'error': 'Sin respuesta de Claude', 'detail': j}), 500

        # ── Gemini ──────────────────────────────────────────
        elif modelo == 'gemini':
            url = f'https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-lite:generateContent?key={GEMINI_KEY}'
            r = req.post(url,
                headers={'Content-Type':'application/json'},
                json={'system_instruction':{'parts':[{'text':sys_prompt}]},
                      'contents':[{'parts':[{'text':user_msg}]}],
                      'generationConfig':{'maxOutputTokens':800}},
                timeout=30)
            j = r.json()
            print(f'[demo/gemini] status={r.status_code} response={json.dumps(j)[:300]}')
            if 'candidates' in j and j['candidates']:
                return jsonify({'html': j['candidates'][0]['content']['parts'][0]['text']})
            # Devolver el error completo al frontend para debug
            error_msg = j.get('error', {}).get('message', 'Sin respuesta de Gemini')
            return jsonify({'error': f'Gemini: {error_msg}', 'detail': j}), 500

        # ── Groq ─────────────────────────────────────────────
        elif modelo == 'groq':
            r = req.post('https://api.groq.com/openai/v1/chat/completions',
                headers={'Content-Type':'application/json','Authorization':f'Bearer {GROQ_KEY}'},
                json={'model':'llama-3.3-70b-versatile','max_tokens':800,
                      'messages':[{'role':'system','content':sys_prompt},{'role':'user','content':user_msg}]},
                timeout=30)
            j = r.json()
            if 'choices' in j and j['choices']:
                return jsonify({'html': j['choices'][0]['message']['content']})
            return jsonify({'error': 'Sin respuesta de Groq', 'detail': j}), 500

        # ── OpenRouter ────────────────────────────────────
        elif modelo == 'openrouter':
            r = req.post('https://openrouter.ai/api/v1/chat/completions',
                headers={'Content-Type':'application/json','Authorization':f'Bearer {OPENROUTER_KEY}'},
                json={'model':'nousresearch/hermes-3-llama-3.1-405b:free','max_tokens':800,
                      'messages':[{'role':'system','content':sys_prompt},{'role':'user','content':user_msg}]},
                timeout=30)
            j = r.json()
            if 'choices' in j and j['choices']:
                return jsonify({'html': j['choices'][0]['message']['content']})
            return jsonify({'error': 'Sin respuesta de OpenRouter', 'detail': j}), 500

        return jsonify({'error': f'Modelo desconocido: {modelo}'}), 400

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/demo/cita', methods=['POST'])
def demo_cita():
    """Recibe solicitud de cita del demo y envía correo a DEMO_MAIL."""
    try:
        d = request.get_json(force=True)

        # Construir cuerpo del correo en HTML
        nombre   = d.get('nombre','')
        cuerpo = f"""
<div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;color:#1a1a18">

<div style="background:#f0f7ff;border-bottom:3px solid #4299e1;padding:24px 28px;border-radius:8px 8px 0 0">
  <h2 style="margin:0;font-size:20px;color:#2d3748">📋 Nueva solicitud de cita</h2>
  <p style="margin:6px 0 0;font-size:13px;color:#718096">Formulario demo · metodo.centrocarvajal.com/demo</p>
</div>

<div style="background:#ffffff;padding:24px 28px;border:1px solid #bee3f8;border-top:none">

  <h3 style="font-size:12px;color:#4299e1;text-transform:uppercase;letter-spacing:.1em;margin:0 0 12px">👤 Datos del paciente</h3>
  <table style="font-size:14px;border-collapse:collapse;width:100%;margin-bottom:24px">
    <tr><td style="padding:5px 0;color:#718096;width:160px">Nombre</td><td style="padding:5px 0"><strong>{nombre}</strong></td></tr>
    <tr><td style="padding:5px 0;color:#718096">Correo</td><td style="padding:5px 0">{d.get('email','')}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">WhatsApp</td><td style="padding:5px 0">{d.get('tel','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Edad</td><td style="padding:5px 0">{d.get('edad','')}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Género</td><td style="padding:5px 0">{d.get('genero','')}</td></tr>
    {f'<tr><td style="padding:5px 0;color:#718096">Embarazo/lactancia</td><td style="padding:5px 0">{d.get("embarazo","")}</td></tr>' if d.get('genero','') == 'mujer' else ''}
    <tr><td style="padding:5px 0;color:#718096">Cómo nos conoció</td><td style="padding:5px 0">{d.get('referido','') or '—'}</td></tr>
  </table>

  <h3 style="font-size:12px;color:#4299e1;text-transform:uppercase;letter-spacing:.1em;margin:0 0 12px">🎯 Perfil estético</h3>
  <table style="font-size:14px;border-collapse:collapse;width:100%;margin-bottom:24px">
    <tr><td style="padding:5px 0;color:#718096;width:160px">Áreas de interés</td><td style="padding:5px 0">{d.get('areas','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Problemas faciales</td><td style="padding:5px 0">{d.get('piel_checks','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Tono de piel</td><td style="padding:5px 0">{d.get('tono','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Zonas corporales</td><td style="padding:5px 0">{d.get('cuerpo_zona','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Objetivos corporales</td><td style="padding:5px 0">{d.get('cuerpo_obj','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Zonas de depilación</td><td style="padding:5px 0">{d.get('vello_zonas','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Situación capilar</td><td style="padding:5px 0">{d.get('capilar','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Tratamientos previos</td><td style="padding:5px 0">{d.get('trat_prev','') or 'Ninguno'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Mayor preocupación</td><td style="padding:5px 0">{d.get('prioridad','') or '—'}</td></tr>
  </table>

  <h3 style="font-size:12px;color:#4299e1;text-transform:uppercase;letter-spacing:.1em;margin:0 0 12px">💡 Hábitos y preferencias</h3>
  <table style="font-size:14px;border-collapse:collapse;width:100%;margin-bottom:24px">
    <tr><td style="padding:5px 0;color:#718096;width:160px">Actividad física</td><td style="padding:5px 0">{d.get('ejercicio','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Nivel de estrés</td><td style="padding:5px 0">{d.get('estres','')}/10</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Urgencia</td><td style="padding:5px 0">{d.get('urgencia','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Presupuesto mensual</td><td style="padding:5px 0">{d.get('presupuesto','') or '—'}</td></tr>
  </table>

  <h3 style="font-size:12px;color:#4299e1;text-transform:uppercase;letter-spacing:.1em;margin:0 0 12px">📅 Cita solicitada</h3>
  <table style="font-size:14px;border-collapse:collapse;width:100%;margin-bottom:8px">
    <tr><td style="padding:5px 0;color:#718096;width:160px">Fecha</td><td style="padding:5px 0"><strong>{d.get('fecha','')}</strong></td></tr>
    <tr><td style="padding:5px 0;color:#718096">Horario</td><td style="padding:5px 0"><strong>{d.get('horario','')}</strong></td></tr>
    <tr><td style="padding:5px 0;color:#718096">Nota adicional</td><td style="padding:5px 0">{d.get('nota','') or '—'}</td></tr>
  </table>

{'<h3 style="font-size:12px;color:#4299e1;text-transform:uppercase;letter-spacing:.1em;margin:20px 0 12px">✨ Evaluación generada por IA</h3><div style="background:#ebf8ff;border:1px solid #bee3f8;border-radius:8px;padding:16px 20px;font-size:14px;color:#2d3748;line-height:1.8">' + d.get('evaluacion_ia','') + '</div>' if d.get('evaluacion_ia') else ''}

<div style="background:#f0f7ff;padding:14px 28px;border:1px solid #bee3f8;border-top:none;border-radius:0 0 8px 8px;text-align:center">
  
</div>

</div>
"""

        resend_payload = {
            'from':    f'Evaluación IA <{MAIL_FROM}>',
            'to':      [DEMO_MAIL],
            'subject': f'Nueva solicitud de cita para — {nombre}',
            'html':    cuerpo
        }

        r = req.post('https://api.resend.com/emails',
            headers={'Authorization': f'Bearer {RESEND_KEY}', 'Content-Type': 'application/json'},
            json=resend_payload, timeout=15)

        print(f'[demo/cita] resend status={r.status_code} body={r.text[:200]}')

        if r.status_code != 200:
            return jsonify({'ok': False, 'detail': r.text}), 500

        # ── Correo al paciente ────────────────────────────────
        email_paciente = d.get('email', '').strip()
        if email_paciente and d.get('evaluacion_ia'):
            cuerpo_paciente = f"""
<div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;color:#2d3748">

<div style="background:linear-gradient(135deg,#fbb6ce,#d6bcfa);padding:28px;border-radius:8px 8px 0 0;text-align:center">
  <h2 style="margin:0;font-size:22px;color:#2d3748">Tu evaluación personalizada ✨</h2>
  <p style="margin:8px 0 0;font-size:14px;color:#4a5568">Hola {nombre}, aquí está tu evaluación estética</p>
</div>

<div style="background:#ffffff;padding:24px 28px;border:1px solid #bee3f8;border-top:none">

  <h3 style="font-size:12px;color:#4299e1;text-transform:uppercase;letter-spacing:.1em;margin:0 0 16px">✨ Recomendaciones para ti</h3>
  <div style="background:#ebf8ff;border:1px solid #bee3f8;border-radius:8px;padding:16px 20px;font-size:14px;color:#2d3748;line-height:1.8;margin-bottom:24px">
    {d.get('evaluacion_ia','')}
  </div>

  <h3 style="font-size:12px;color:#4299e1;text-transform:uppercase;letter-spacing:.1em;margin:0 0 12px">📅 Tu cita solicitada</h3>
  <table style="font-size:14px;border-collapse:collapse;width:100%;margin-bottom:8px">
    <tr><td style="padding:5px 0;color:#718096;width:140px">Fecha</td><td style="padding:5px 0"><strong>{d.get('fecha','')}</strong></td></tr>
    <tr><td style="padding:5px 0;color:#718096">Horario</td><td style="padding:5px 0"><strong>{d.get('horario','')}</strong></td></tr>
    {f'<tr><td style="padding:5px 0;color:#718096">Nota</td><td style="padding:5px 0">{d.get("nota","")}</td></tr>' if d.get('nota') else ''}
  </table>

  <p style="font-size:13px;color:#718096;margin-top:20px;line-height:1.7">
    Nos pondremos en contacto contigo pronto para confirmar tu cita. Si tienes alguna pregunta, no dudes en escribirnos.
  </p>

</div>

<div style="background:#f0f7ff;padding:14px 28px;border:1px solid #bee3f8;border-top:none;border-radius:0 0 8px 8px;text-align:center">
  <p style="font-size:11px;color:#718096;margin:0">Tu Estética · Evaluación personalizada con IA</p>
</div>

</div>
"""
            resend_paciente = {
                'from':    f'Tu Estética <{MAIL_FROM}>',
                'to':      [email_paciente],
                'subject': f'{nombre}, aquí está tu evaluación estética personalizada ✨',
                'html':    cuerpo_paciente
            }
            rp = req.post('https://api.resend.com/emails',
                headers={'Authorization': f'Bearer {RESEND_KEY}', 'Content-Type': 'application/json'},
                json=resend_paciente, timeout=15)
            print(f'[demo/cita] paciente email status={rp.status_code}')

        return jsonify({'ok': True})

    except Exception as e:
        print(f'[demo/cita] error: {e}')
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/formulario', methods=['GET'])
def formulario():
    # Formulario web completo para el paciente
    with open(os.path.join(os.path.dirname(__file__), 'formulario.html'), encoding='utf-8') as f:
        return f.read()

@app.route('/planes_generados/<path:filename>')
def serve_plan(filename):
    return send_from_directory(PLANES_DIR, filename)

@app.route('/status')
def status():
    job_id = request.args.get('job', '')
    job = jobs.get(job_id, {'status': 'working', 'msg': 'Procesando...', 'pct': 5})
    return jsonify(job)


# ── Endpoint formulario web (/enviar) ─────────────────────────
@app.route('/enviar', methods=['POST'])
def enviar():
    raw = request.form.get('data', '')
    if not raw:
        return jsonify({'error': 'No se recibieron datos del formulario'}), 400

    try:
        form = json.loads(raw)
    except Exception:
        return jsonify({'error': 'JSON invalido en los datos del formulario'}), 400

    data = _mapear_formulario(form)

    if not data.get('nombre'):
        return jsonify({'error': 'El campo Nombre completo es obligatorio'}), 400

    # Guardar fotos opcionales (se envían al correo del staff)
    fotos = []
    for i in range(1, 5):
        foto = request.files.get(f'foto_{i}')
        if foto:
            tmp = f'/tmp/foto_{uuid.uuid4().hex}_{foto.filename}'
            foto.save(tmp)
            fotos.append(tmp)

    job_id = uuid.uuid4().hex[:16]
    jobs[job_id] = {'status': 'working', 'msg': 'Iniciando generacion del plan...'}

    modelo = request.form.get('modelo', 'claude')  # viene del FormData, no del JSON
    t = threading.Thread(target=worker, args=(job_id, data, [], fotos, modelo), daemon=True)
    t.start()

    return jsonify({'jobId': job_id, 'nombre': data['nombre'], 'modelo': modelo})


# ── Endpoint carga .docx (/upload) ────────────────────────────
@app.route('/upload', methods=['POST'])
def upload():
    if 'docx' not in request.files:
        return jsonify({'error': 'No se recibio archivo'}), 400

    f = request.files['docx']
    tmp_path = f'/tmp/carvajal_{uuid.uuid4().hex}.docx'
    f.save(tmp_path)

    texto = leer_docx(tmp_path)
    os.unlink(tmp_path)

    if not texto:
        return jsonify({'error': 'No se pudo leer el archivo .docx'}), 400

    data, faltantes = parsear_cuestionario(texto)
    if not data.get('nombre'):
        return jsonify({'error': 'No se encontro nombre_completo en el documento'}), 400

    job_id = uuid.uuid4().hex[:16]
    jobs[job_id] = {'status': 'working', 'msg': 'Iniciando generacion del plan...'}

    modelo = request.form.get('modelo', 'claude')
    t = threading.Thread(target=worker, args=(job_id, data, faltantes, [], modelo), daemon=True)
    t.start()

    return jsonify({'jobId': job_id, 'nombre': data['nombre'], 'modelo': modelo})


# ════════════════════════════════════════════════════════════
# RUTA BORRADOR EDITABLE
# ════════════════════════════════════════════════════════════

@app.route('/borrador/<job_id>', methods=['GET'])
def ver_borrador(job_id):
    """Sirve el borrador editable. Descarga desde Cloudinary si no está en memoria."""
    from flask import Response
    html = descargar_borrador_cloudinary(job_id)
    if html:
        return Response(html, content_type='text/html; charset=utf-8')
    job = jobs.get(job_id)
    if not job or job.get('status') != 'done':
        return Response('<h2 style="font-family:sans-serif;padding:40px;color:#666">Borrador no encontrado o aún procesando.</h2>', status=404, content_type='text/html; charset=utf-8')
    return Response('<h2 style="font-family:sans-serif;padding:40px;color:#666">Borrador no disponible. Verifica Cloudinary.</h2>', status=404, content_type='text/html; charset=utf-8')


@app.route('/guardar/<job_id>', methods=['POST'])
def guardar_borrador(job_id):
    """Recibe el HTML completo editado y lo guarda en Cloudinary."""
    from flask import Response as R
    html = request.get_data(as_text=True)
    if not html:
        return jsonify({'ok': False, 'error': 'HTML vacío'}), 400
    url = subir_borrador_cloudinary(html, job_id)
    # Actualizar nombre en jobs si existe
    if job_id in jobs:
        jobs[job_id]['borrador_actualizado'] = True
    return jsonify({'ok': True, 'url': url or ''})





# ════════════════════════════════════════════════════════════
# WORKER
# ════════════════════════════════════════════════════════════

def subir_plan_cloudinary(html_path, html_name, job_id=''):
    """Sube el HTML del plan a Cloudinary como raw file.
    Devuelve la URL pública o None si falla."""
    if not CLOUDINARY_CLOUD_NAME:
        print('Cloudinary no configurado, usando URL local')
        return None
    try:
        resultado = cloudinary.uploader.upload(
            html_path,
            folder        = 'carvajal/planes',
            public_id     = re.sub(r'\.html$', '', html_name),
            resource_type = 'raw',
            overwrite     = True,
            context       = f'job_id={job_id}' if job_id else None,
        )
        url = resultado.get('secure_url', '')
        print(f'Cloudinary OK: {url[:80]}')
        return url
    except Exception as e:
        print(f'Cloudinary error: {e}')
        return None


def subir_borrador_cloudinary(html_content, job_id):
    """Sube el HTML del borrador a Cloudinary como raw file."""
    if not CLOUDINARY_CLOUD_NAME:
        return None
    try:
        import tempfile
        with tempfile.NamedTemporaryFile(mode='wb', suffix='.html', delete=False) as tmp:
            tmp.write(html_content.encode('utf-8'))  # escribir bytes UTF-8 explícito
            tmp_path = tmp.name
        resultado = cloudinary.uploader.upload(
            tmp_path,
            folder='carvajal/borradores',
            public_id=f'borrador_{job_id}',
            resource_type='raw',
            overwrite=True,
        )
        os.unlink(tmp_path)
        return resultado.get('secure_url', '')
    except Exception as e:
        print(f'Error subiendo borrador: {e}')
        return None


def descargar_borrador_cloudinary(job_id):
    """Descarga el HTML del borrador desde Cloudinary."""
    if not CLOUDINARY_CLOUD_NAME:
        return None
    try:
        url = f'https://res.cloudinary.com/{CLOUDINARY_CLOUD_NAME}/raw/upload/carvajal/borradores/borrador_{job_id}.html'
        r = req.get(url, timeout=30)
        if r.status_code == 200:
            return r.content.decode('utf-8')  # forzar UTF-8, Cloudinary no declara charset
        return None
    except Exception as e:
        print(f'Error descargando borrador: {e}')
        return None


def render_borrador(plan_json, data, job_id):
    """Genera el HTML del borrador — ahora usa la misma plantilla que render_plan."""
    return render_plan(plan_json, data, job_id=job_id)


def _render_borrador_legacy(plan_json, data, job_id):
    """LEGACY — ya no se usa. render_borrador ahora llama render_plan."""
    tpl_path = os.path.join(os.path.dirname(__file__), 'plantilla_borrador.html')
    with open(tpl_path, encoding='utf-8') as f:
        tpl = f.read()

    def esc(s): return str(s).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')

    nombre = data.get('nombre', '')

    # Diagnóstico rows
    badge_map = {'warning': 'badge-warning', 'critical': 'badge-critical', 'normal': 'badge-normal'}
    badge_label = {'warning': '⚠ Atención', 'critical': '✕ Crítico', 'normal': 'Normal'}
    diag_html = ''.join(
        f'<tr><td style="font-weight:500">{esc(f.get("area",""))}</td><td contenteditable="true">{esc(f.get("estado",""))}</td><td contenteditable="true">{esc(f.get("hallazgos",""))}</td><td><span class="{badge_map.get(f.get("alerta","normal"),"badge-normal")}">{badge_label.get(f.get("alerta","normal"),"Normal")}</span></td></tr>'
        for f in plan_json.get('diagnostico', {}).get('filas', [])
    )

    # Rutina
    tag_map = {'Nutricion':'tag-n','Sueno':'tag-s','Actividad':'tag-a','Mental':'tag-m','Estetico':'tag-e','Salud':'tag-h'}
    rutina_html = ''.join(
        f'<div class="rutina-row"><span class="rutina-hora">{esc(r["hora"])}</span><span contenteditable="true" style="flex:1">{esc(r["actividad"])}</span><span class="rutina-tag {tag_map.get(r["pilar"],"tag-n")}">{esc(r["pilar"])}</span></div>'
        for r in plan_json.get('rutina', {}).get('items', [])
    )

    # Pilares helpers
    p1 = plan_json.get('pilar1', {})
    p2 = plan_json.get('pilar2', {})
    p3 = plan_json.get('pilar3', {})
    p4 = plan_json.get('pilar4', {})
    p5 = plan_json.get('pilar5', {})
    comp = plan_json.get('compromiso', {})

    p1_perm = ''.join(f'<li contenteditable="true">{esc(i)}</li>' for i in p1.get('permitidos',[]))
    p1_evit = ''.join(f'<li contenteditable="true">{esc(i)}</li>' for i in p1.get('evitar',[]))
    p1_menu = ''.join(
        f'<tr><td class="dia">{esc(m.get("dia",""))}</td><td contenteditable="true">{esc(m.get("desayuno",""))}</td><td contenteditable="true">{esc(m.get("almuerzo",""))}</td><td contenteditable="true">{esc(m.get("cena",""))}</td><td contenteditable="true">{esc(m.get("snack",""))}</td></tr>'
        for m in p1.get('menu',[])
    )
    p1_supl = ''
    if p1.get('suplementacion'):
        items = ''.join(f'<li contenteditable="true" style="font-size:13px;padding:4px 0;border-bottom:1px solid #e5e7eb">{esc(s)}</li>' for s in p1['suplementacion'])
        p1_supl = f'<h4 style="font-size:12px;font-weight:500;color:#6b7280;text-transform:uppercase;letter-spacing:.06em;margin-bottom:8px">Suplementación</h4><ul style="list-style:none">{items}</ul>'

    p3_tec = ''.join(f'<li contenteditable="true">{esc(t)}</li>' for t in p3.get('tecnicas',[]))

    p4_proto = ''.join(f'<li contenteditable="true" style="padding:4px 0;font-size:13px">{esc(s)}</li>' for s in p4.get('protocolo',[]))
    p4_reg = ''.join(f'<li contenteditable="true">{esc(r)}</li>' for r in p4.get('reglas',[]))

    # Pilar 5 bimestres
    p5_bim = ''
    for bim in p5.get('bimestres',[]):
        rows = ''.join(f'<tr><td contenteditable="true" style="font-weight:500">{esc(t.get("nombre",""))}</td><td contenteditable="true">{esc(t.get("sesiones",""))}</td><td contenteditable="true" style="font-weight:500">{esc(t.get("inversion",""))}</td><td contenteditable="true">{esc(t.get("beneficio",""))}</td></tr>' for t in bim.get('tratamientos',[]))
        p5_bim += f'<div style="margin-bottom:16px"><div style="font-size:12px;font-weight:600;color:#8fa832;text-transform:uppercase;letter-spacing:.06em;margin-bottom:6px" contenteditable="true">{esc(bim.get("periodo",""))} · {esc(bim.get("titulo",""))}</div><table><thead><tr><th>Tratamiento</th><th>Sesiones</th><th>Inversión</th><th>Beneficio</th></tr></thead><tbody>{rows}</tbody></table><div style="font-size:12px;color:#6b7280;margin-top:6px;text-align:right" contenteditable="true">Total bimestre: ${bim.get("total",0):,}</div></div>'

    notas = p5.get('notas_criticas',[])
    p5_notas = ''
    if notas:
        items = ''.join(f'<p style="font-size:13px;padding:3px 0;border-bottom:1px solid #e5e7eb" contenteditable="true">{esc(n)}</p>' for n in notas)
        p5_notas = f'<div style="background:#fff7ed;border:1px solid #fed7aa;border-radius:4px;padding:12px 16px;margin-bottom:12px"><strong style="font-size:11px;color:#9a3412;text-transform:uppercase;letter-spacing:.06em;display:block;margin-bottom:6px">Notas críticas</strong>{items}</div>'

    p5_am = ''.join(f'<div style="display:flex;gap:8px;align-items:flex-start;padding:6px 0;border-bottom:1px solid #e5e7eb"><span style="font-size:11px;font-weight:600;background:#8fa832;color:#fff;border-radius:50%;width:18px;height:18px;display:flex;align-items:center;justify-content:center;flex-shrink:0">{s.get("paso","")}</span><div><div style="font-size:12px;font-weight:500" contenteditable="true">{esc(s.get("producto",""))}</div><div style="font-size:11px;color:#6b7280" contenteditable="true">{esc(s.get("descripcion",""))}</div></div></div>' for s in p5.get('rutina_am',[]))
    p5_pm = ''.join(f'<div style="display:flex;gap:8px;align-items:flex-start;padding:6px 0;border-bottom:1px solid #e5e7eb"><span style="font-size:11px;font-weight:600;background:#2d3a2e;color:#fff;border-radius:50%;width:18px;height:18px;display:flex;align-items:center;justify-content:center;flex-shrink:0">{s.get("paso","")}</span><div><div style="font-size:12px;font-weight:500" contenteditable="true">{esc(s.get("producto",""))}</div><div style="font-size:11px;color:#6b7280" contenteditable="true">{esc(s.get("descripcion",""))}</div></div></div>' for s in p5.get('rutina_pm',[]))

    comp_res = ''.join(f'<li contenteditable="true">{esc(r["texto"])}</li>' for r in comp.get('resultados',[]))
    comp_pasos = ''.join(f'<li style="padding:4px 0;font-size:13px" contenteditable="true">{esc(p)}</li>' for p in comp.get('proximos_pasos',[]))

    replacements = {
        '{{JOB_ID}}': job_id,
        '{{NOMBRE}}': esc(nombre),
        '{{EDAD}}': esc(data.get('edad','')),
        '{{OCUPACION}}': esc(data.get('ocupacion','')),
        '{{FECHA}}': esc(data.get('fecha','')),
        '{{DIAGNOSTICO_FILAS}}': diag_html,
        '{{RUTINA_NOTA}}': esc(plan_json.get('rutina',{}).get('nota','')),
        '{{RUTINA_FILAS}}': rutina_html,
        '{{P1_TITULO}}': esc(p1.get('titulo','')),
        '{{P1_OBJETIVO}}': esc(p1.get('objetivo','')),
        '{{P1_PERMITIDOS}}': p1_perm,
        '{{P1_EVITAR}}': p1_evit,
        '{{P1_MENU}}': p1_menu,
        '{{P1_SUPLEMENTACION}}': p1_supl,
        '{{P2_TITULO}}': esc(p2.get('titulo','')),
        '{{P2_OBJETIVO}}': esc(p2.get('objetivo','')),
        '{{P2_PLAN}}': esc(p2.get('plan_semanal','')),
        '{{P2_ADAPTACIONES}}': esc(p2.get('adaptaciones','')),
        '{{P3_TITULO}}': esc(p3.get('titulo','')),
        '{{P3_OBJETIVO}}': esc(p3.get('objetivo','')),
        '{{P3_TECNICAS}}': p3_tec,
        '{{P4_TITULO}}': esc(p4.get('titulo','')),
        '{{P4_OBJETIVO}}': esc(p4.get('objetivo','')),
        '{{P4_PROTOCOLO}}': p4_proto,
        '{{P4_REGLAS}}': p4_reg,
        '{{P5_TITULO}}': esc(p5.get('titulo','')),
        '{{P5_OBJETIVO}}': esc(p5.get('objetivo','')),
        '{{P5_BIMESTRES}}': p5_bim,
        '{{P5_NOTAS_CRITICAS}}': p5_notas,
        '{{P5_RUTINA_AM}}': p5_am,
        '{{P5_RUTINA_PM}}': p5_pm,
        '{{COMP_PARRAFO}}': esc(comp.get('parrafo','')),
        '{{COMP_RESULTADOS}}': comp_res,
        '{{COMP_PASOS}}': comp_pasos,
    }
    for k, v in replacements.items():
        tpl = tpl.replace(k, v)
    return tpl



# ════════════════════════════════════════════════════════════
# ADMIN — Página /planes con login
# ════════════════════════════════════════════════════════════

@app.route('/planes', methods=['GET'])
def admin_planes():
    return PLANES_HTML

PLANES_HTML = """<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Centro Carvajal · Planes Generados</title>
<style>
:root{
  --dark:#1a1410;--olive:#8fa832;--olive-light:rgba(143,168,50,.1);
  --gold:#b8935a;--cream:#f4f5ef;--white:#fff;--gray:#6b7280;
  --border:rgba(143,168,50,.2);--red:#dc2626;
}
*{margin:0;padding:0;box-sizing:border-box}
body{font-family:'Segoe UI',system-ui,sans-serif;background:var(--cream);color:var(--dark);min-height:100vh}

/* ── LOGIN ── */
#login-screen{
  display:flex;align-items:center;justify-content:center;
  min-height:100vh;padding:20px;
}
.login-box{
  background:var(--white);border-radius:12px;padding:44px 40px;
  max-width:400px;width:100%;box-shadow:0 8px 40px rgba(0,0,0,.1);
  text-align:center;
}
.login-logo{margin:0 auto 20px}
.login-box h2{font-size:22px;font-weight:600;color:var(--dark);margin-bottom:6px}
.login-box p{font-size:13px;color:var(--gray);margin-bottom:28px}
.login-field{
  width:100%;padding:13px 16px;border:1px solid #d4dcc0;
  border-radius:8px;font-size:14px;margin-bottom:14px;
  font-family:inherit;outline:none;transition:border-color .2s;
}
.login-field:focus{border-color:var(--olive)}
.login-btn{
  width:100%;padding:14px;background:var(--dark);color:var(--cream);
  border:none;border-radius:8px;font-size:14px;font-weight:500;
  cursor:pointer;font-family:inherit;transition:background .2s;
}
.login-btn:hover{background:#2d3a2e}
.login-error{
  display:none;background:#fef2f2;border:1px solid #fecaca;
  border-radius:6px;padding:10px 14px;font-size:13px;color:var(--red);
  margin-bottom:14px;text-align:left;
}

/* ── APP SHELL ── */
#app{display:none;min-height:100vh;flex-direction:column}

/* Topbar */
.topbar{
  background:var(--dark);padding:14px 28px;
  display:flex;align-items:center;justify-content:space-between;
  position:sticky;top:0;z-index:100;
}
.topbar-left{display:flex;align-items:center;gap:14px}
.topbar-logo{
  width:36px;height:36px;border-radius:50%;
  border:1.5px solid var(--gold);
  display:flex;align-items:center;justify-content:center;
  font-family:Georgia,serif;font-size:12px;color:var(--gold);font-weight:600;
}
.topbar-title{
  font-size:15px;font-weight:600;color:var(--white);
}
.topbar-sub{font-size:11px;color:rgba(255,255,255,.35);margin-top:1px}
.topbar-right{display:flex;align-items:center;gap:12px}
.topbar-count{
  background:var(--olive-light);border:1px solid var(--border);
  border-radius:20px;padding:4px 12px;
  font-size:12px;font-weight:600;color:var(--olive);
}
.logout-btn{
  background:transparent;border:1px solid rgba(255,255,255,.15);
  color:rgba(255,255,255,.5);border-radius:6px;padding:6px 12px;
  font-size:12px;cursor:pointer;font-family:inherit;transition:all .2s;
}
.logout-btn:hover{border-color:rgba(255,255,255,.4);color:var(--white)}

/* Main content */
.main{flex:1;padding:28px;max-width:1200px;margin:0 auto;width:100%}

/* Toolbar */
.toolbar{
  display:flex;align-items:center;gap:12px;margin-bottom:24px;flex-wrap:wrap;
}
.search-wrap{position:relative;flex:1;min-width:220px}
.search-input{
  width:100%;padding:10px 16px 10px 38px;
  border:1px solid #d4dcc0;border-radius:8px;font-size:13px;
  font-family:inherit;outline:none;background:var(--white);
  transition:border-color .2s;
}
.search-input:focus{border-color:var(--olive)}
.search-icon{
  position:absolute;left:12px;top:50%;transform:translateY(-50%);
  color:var(--gray);font-size:15px;pointer-events:none;
}
.filter-select{
  padding:10px 14px;border:1px solid #d4dcc0;border-radius:8px;
  font-size:13px;font-family:inherit;background:var(--white);
  outline:none;cursor:pointer;color:var(--dark);
}
.refresh-btn{
  padding:10px 16px;background:var(--white);border:1px solid #d4dcc0;
  border-radius:8px;font-size:13px;cursor:pointer;font-family:inherit;
  color:var(--dark);transition:all .2s;display:flex;align-items:center;gap:6px;
}
.refresh-btn:hover{border-color:var(--olive);color:var(--olive)}
.refresh-btn.spinning svg{animation:spin .8s linear infinite}
@keyframes spin{to{transform:rotate(360deg)}}

/* Stats bar */
.stats{
  display:grid;grid-template-columns:repeat(4,1fr);gap:14px;margin-bottom:24px;
}
.stat-card{
  background:var(--white);border:1px solid rgba(0,0,0,.06);
  border-radius:10px;padding:16px 20px;
}
.stat-val{font-size:24px;font-weight:700;color:var(--dark);margin-bottom:2px}
.stat-label{font-size:11px;color:var(--gray);text-transform:uppercase;letter-spacing:.8px}
.stat-card.olive .stat-val{color:var(--olive)}
.stat-card.gold .stat-val{color:var(--gold)}

/* Plans grid */
.plans-grid{
  display:grid;
  grid-template-columns:repeat(auto-fill,minmax(320px,1fr));
  gap:16px;
}

/* Plan card */
.plan-card{
  background:var(--white);border:1px solid rgba(0,0,0,.07);
  border-radius:10px;overflow:hidden;transition:box-shadow .2s,transform .15s;
}
.plan-card:hover{box-shadow:0 4px 20px rgba(0,0,0,.1);transform:translateY(-2px)}

.card-header{
  background:var(--dark);padding:14px 18px;
  display:flex;align-items:flex-start;justify-content:space-between;gap:10px;
}
.card-nombre{
  font-size:15px;font-weight:600;color:var(--white);
  line-height:1.3;flex:1;
}
.card-modelo{
  font-size:10px;font-weight:700;padding:3px 8px;border-radius:12px;
  text-transform:uppercase;letter-spacing:.5px;flex-shrink:0;margin-top:2px;
}
.modelo-claude{background:rgba(255,255,255,.15);color:rgba(255,255,255,.8)}
.modelo-gemini{background:rgba(26,115,232,.3);color:#93c5fd}
.modelo-groq{background:rgba(245,80,54,.3);color:#fca5a5}
.modelo-otro{background:rgba(255,255,255,.1);color:rgba(255,255,255,.6)}

.card-body{padding:16px 18px}
.card-meta{display:flex;flex-direction:column;gap:5px;margin-bottom:14px}
.card-meta-row{
  display:flex;align-items:center;gap:7px;
  font-size:12px;color:var(--gray);
}
.card-meta-row svg{width:13px;height:13px;stroke:currentColor;fill:none;stroke-width:1.8;flex-shrink:0}
.card-meta-row strong{color:var(--dark)}

.card-actions{display:flex;gap:8px}
.btn-ver{
  flex:1;padding:9px;background:var(--olive-light);
  border:1px solid var(--border);border-radius:6px;
  font-size:12px;font-weight:600;color:var(--olive);
  text-decoration:none;text-align:center;
  transition:all .2s;cursor:pointer;
  display:flex;align-items:center;justify-content:center;gap:5px;
}
.btn-ver:hover{background:var(--olive);color:var(--white);border-color:var(--olive)}
.btn-editar{
  flex:1;padding:9px;background:#f0f9e8;
  border:1px solid rgba(143,168,50,.3);border-radius:6px;
  font-size:12px;font-weight:600;color:#5a7a1a;
  text-decoration:none;text-align:center;
  transition:all .2s;cursor:pointer;
  display:flex;align-items:center;justify-content:center;gap:5px;
}
.btn-editar:hover{background:#2d3a2e;color:var(--white);border-color:#2d3a2e}

/* Empty / loading */
.state-box{
  text-align:center;padding:80px 20px;color:var(--gray);
  grid-column:1/-1;
}
.state-box .icon{font-size:48px;margin-bottom:16px}
.state-box h3{font-size:18px;color:var(--dark);margin-bottom:8px}
.state-box p{font-size:13px;line-height:1.6;max-width:340px;margin:0 auto}

/* Skeleton */
.skeleton{
  background:linear-gradient(90deg,#f0f0f0 25%,#e0e0e0 50%,#f0f0f0 75%);
  background-size:200% 100%;
  animation:shimmer 1.5s infinite;
  border-radius:6px;
}
@keyframes shimmer{0%{background-position:200% 0}100%{background-position:-200% 0}}
.skel-card{
  background:var(--white);border:1px solid rgba(0,0,0,.07);
  border-radius:10px;overflow:hidden;
}
.skel-header{height:68px;background:#2d2d2d}
.skel-body{padding:16px 18px}
.skel-line{height:12px;margin-bottom:8px}
.skel-actions{display:flex;gap:8px;margin-top:14px}
.skel-btn{height:36px;flex:1;border-radius:6px}

/* Toast */
.toast{
  position:fixed;bottom:24px;right:24px;
  background:var(--dark);color:var(--white);
  padding:12px 20px;border-radius:8px;font-size:13px;
  box-shadow:0 4px 20px rgba(0,0,0,.2);
  transform:translateY(100px);opacity:0;
  transition:all .3s;z-index:999;
  display:flex;align-items:center;gap:8px;
}
.toast.show{transform:translateY(0);opacity:1}
.toast.success{border-left:3px solid var(--olive)}
.toast.error{border-left:3px solid var(--red)}
</style>
</head>
<body>

<!-- ══ LOGIN ══ -->
<div id="login-screen" style="display:none">
  <div class="login-box">
    <svg class="login-logo" width="48" height="48" viewBox="0 0 48 48" fill="none">
      <circle cx="24" cy="24" r="22" stroke="#b8935a" stroke-width="1.5"/>
      <text x="24" y="29" text-anchor="middle" fill="#b8935a" font-family="Georgia,serif" font-size="14" font-weight="600">CC</text>
    </svg>
    <h2>Centro Carvajal</h2>
    <p>Panel de administración · Planes generados</p>
    <div class="login-error" id="login-error">Contraseña incorrecta. Inténtalo de nuevo.</div>
    <input type="password" class="login-field" id="login-pass" placeholder="Contraseña" onkeydown="if(event.key==='Enter')doLogin()">
    <button class="login-btn" onclick="doLogin()">Ingresar</button>
  </div>
</div>

<!-- ══ APP ══ -->
<div id="app" style="display:flex">

  <!-- Topbar -->
  <div class="topbar">
    <div class="topbar-left">
      <div class="topbar-logo">CC</div>
      <div>
        <div class="topbar-title">Planes Generados</div>
        <div class="topbar-sub">Centro Carvajal · Panel de Administración</div>
      </div>
    </div>
    <div class="topbar-right">
      <div class="topbar-count" id="total-count">— planes</div>

    </div>
  </div>

  <!-- Main -->
  <div class="main">

    <!-- Stats -->
    <div class="stats" id="stats-bar" style="display:none">
      <div class="stat-card">
        <div class="stat-val" id="stat-total">—</div>
        <div class="stat-label">Planes totales</div>
      </div>
      <div class="stat-card olive">
        <div class="stat-val" id="stat-mes">—</div>
        <div class="stat-label">Este mes</div>
      </div>
      <div class="stat-card gold">
        <div class="stat-val" id="stat-claude">—</div>
        <div class="stat-label">Con Claude</div>
      </div>
      <div class="stat-card">
        <div class="stat-val" id="stat-otros">—</div>
        <div class="stat-label">Gemini / Groq</div>
      </div>
    </div>

    <!-- Toolbar -->
    <div class="toolbar">
      <div class="search-wrap">
        <span class="search-icon">🔍</span>
        <input type="text" class="search-input" id="search-input" placeholder="Buscar por nombre de paciente..." oninput="filtrarPlanes()">
      </div>
      <select class="filter-select" id="filter-modelo" onchange="filtrarPlanes()">
        <option value="">Todos los modelos</option>
        <option value="claude">Claude</option>
        <option value="gemini">Gemini</option>
        <option value="groq">Groq</option>
      </select>
      <select class="filter-select" id="filter-orden" onchange="filtrarPlanes()">
        <option value="desc">Más recientes primero</option>
        <option value="asc">Más antiguos primero</option>
      </select>
      <button class="refresh-btn" id="refresh-btn" onclick="cargarPlanes()">
        <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polyline points="23 4 23 10 17 10"/><path d="M20.49 15a9 9 0 1 1-2.12-9.36L23 10"/></svg>
        Actualizar
      </button>
    </div>

    <!-- Grid -->
    <div class="plans-grid" id="plans-grid">
      <!-- Skeletons mientras carga -->
      <div class="skel-card"><div class="skel-header"></div><div class="skel-body"><div class="skeleton skel-line" style="width:70%"></div><div class="skeleton skel-line" style="width:50%"></div><div class="skeleton skel-line" style="width:60%"></div><div class="skel-actions"><div class="skeleton skel-btn"></div><div class="skeleton skel-btn"></div></div></div></div>
      <div class="skel-card"><div class="skel-header"></div><div class="skel-body"><div class="skeleton skel-line" style="width:65%"></div><div class="skeleton skel-line" style="width:45%"></div><div class="skeleton skel-line" style="width:55%"></div><div class="skel-actions"><div class="skeleton skel-btn"></div><div class="skeleton skel-btn"></div></div></div></div>
      <div class="skel-card"><div class="skel-header"></div><div class="skel-body"><div class="skeleton skel-line" style="width:75%"></div><div class="skeleton skel-line" style="width:55%"></div><div class="skeleton skel-line" style="width:50%"></div><div class="skel-actions"><div class="skeleton skel-btn"></div><div class="skeleton skel-btn"></div></div></div></div>
    </div>
  </div>
</div>


  <!-- ── SECCIÓN CALENDARIO ── -->
  <div style="background:#f4f5ef;border-top:1px solid rgba(143,168,50,.2);padding:28px;margin-top:8px">
    <div style="max-width:1200px;margin:0 auto">
      <div style="margin-bottom:16px">
        <div style="font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:2px;color:var(--olive);margin-bottom:4px">Herramienta</div>
        <div style="font-size:18px;font-weight:700;color:var(--dark)">Generador de Calendario</div>
        <div style="font-size:12px;color:var(--gray);margin-top:3px">Genera un calendario de seguimiento listo para imprimir como PDF desde Chrome</div>
      </div>
      <div style="background:white;border:1px solid rgba(143,168,50,.2);border-radius:10px;padding:20px;display:flex;align-items:flex-end;gap:16px;flex-wrap:wrap">
        <div style="display:flex;flex-direction:column;gap:5px">
          <label style="font-size:11px;font-weight:600;text-transform:uppercase;letter-spacing:1px;color:var(--gray)">Mes de inicio</label>
          <input type="month" id="cal-desde" style="padding:9px 12px;border:1px solid #d4dcc0;border-radius:6px;font-size:13px;font-family:inherit;outline:none;color:var(--dark);min-width:160px">
        </div>
        <div style="display:flex;flex-direction:column;gap:5px">
          <label style="font-size:11px;font-weight:600;text-transform:uppercase;letter-spacing:1px;color:var(--gray)">Mes final</label>
          <input type="month" id="cal-hasta" style="padding:9px 12px;border:1px solid #d4dcc0;border-radius:6px;font-size:13px;font-family:inherit;outline:none;color:var(--dark);min-width:160px">
        </div>
        <button onclick="generarCalendario()" style="padding:10px 24px;background:var(--olive);color:white;border:none;border-radius:6px;font-size:13px;font-weight:600;cursor:pointer;font-family:inherit;display:flex;align-items:center;gap:8px;white-space:nowrap">
          Generar e Imprimir
        </button>
      </div>
      <div style="margin-top:10px;font-size:11px;color:var(--gray)">
        Al generar se abrira una nueva ventana lista para imprimir. Usa <strong>Ctrl+P</strong> &rarr; <strong>Guardar como PDF</strong> &rarr; Orientacion <strong>Horizontal</strong> &rarr; Sin margenes.
      </div>
    </div>
  </div>


<!-- Toast -->
<div class="toast" id="toast"></div>

</div>
<script>

// ── Sin autenticación — acceso libre ────────────────────────
let token = 'open';

function doLogout() {}  // no-op

// ── Datos ────────────────────────────────────────────────────
let todosLosPlanes = [];

async function cargarPlanes() {
  const btn = document.getElementById('refresh-btn');
  btn.classList.add('spinning');

  try {
    const r    = await fetch('/api/planes', {
      headers: {'X-Token': token}
    });
    const data = await r.json();
    todosLosPlanes = data.planes || [];
    actualizarStats();
    filtrarPlanes();
    document.getElementById('stats-bar').style.display = 'grid';
  } catch(e) {
    mostrarToast('Error al cargar planes', 'error');
  } finally {
    btn.classList.remove('spinning');
  }
}

function actualizarStats() {
  const total  = todosLosPlanes.length;
  const ahora  = new Date();
  const mes    = todosLosPlanes.filter(p => {
    if (!p.fecha_raw) return false;
    const f = new Date(p.fecha_raw);
    return f.getMonth() === ahora.getMonth() && f.getFullYear() === ahora.getFullYear();
  }).length;
  const claude = todosLosPlanes.filter(p => p.modelo === 'claude').length;
  const otros  = total - claude;

  document.getElementById('stat-total').textContent  = total;
  document.getElementById('stat-mes').textContent     = mes;
  document.getElementById('stat-claude').textContent  = claude;
  document.getElementById('stat-otros').textContent   = otros;
  document.getElementById('total-count').textContent  = `${total} plan${total !== 1 ? 'es' : ''}`;
}

function filtrarPlanes() {
  const q      = document.getElementById('search-input').value.toLowerCase();
  const modelo = document.getElementById('filter-modelo').value;
  const orden  = document.getElementById('filter-orden').value;

  let lista = todosLosPlanes.filter(p => {
    const matchQ = !q || p.nombre.toLowerCase().includes(q);
    const matchM = !modelo || p.modelo === modelo;
    return matchQ && matchM;
  });

  if (orden === 'asc') lista = lista.slice().reverse();

  renderPlanes(lista);
}

function renderPlanes(lista) {
  const grid = document.getElementById('plans-grid');

  if (!lista.length) {
    grid.innerHTML = `<div class="state-box">
      <div class="icon">📋</div>
      <h3>No hay planes</h3>
      <p>No se encontraron planes con los filtros actuales.</p>
    </div>`;
    return;
  }

  grid.innerHTML = lista.map(p => {
    const modeloClass = {claude:'modelo-claude', gemini:'modelo-gemini', groq:'modelo-groq'}[p.modelo] || 'modelo-otro';
    const modeloLabel = (p.modelo || 'IA').charAt(0).toUpperCase() + (p.modelo || 'ia').slice(1);
    const borrador_url = `/borrador/${p.job_id}`;
    return `<div class="plan-card">
      <div class="card-header">
        <div class="card-nombre">${esc(p.nombre)}</div>
        <span class="card-modelo ${modeloClass}">${modeloLabel}</span>
      </div>
      <div class="card-body">
        <div class="card-meta">
          <div class="card-meta-row">
            <svg viewBox="0 0 24 24"><rect x="3" y="4" width="18" height="18" rx="2"/><line x1="16" y1="2" x2="16" y2="6"/><line x1="8" y1="2" x2="8" y2="6"/><line x1="3" y1="10" x2="21" y2="10"/></svg>
            <span>${esc(p.fecha)}</span>
          </div>
          <div class="card-meta-row">
            <svg viewBox="0 0 24 24"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/></svg>
            <span style="word-break:break-all;font-size:11px;color:#aaa">${esc(p.nombre_archivo)}</span>
          </div>
        </div>
        <div class="card-actions">
          <a href="${esc(p.url)}" target="_blank" class="btn-ver">
            <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M1 12s4-8 11-8 11 8 11 8-4 8-11 8-11-8-11-8z"/><circle cx="12" cy="12" r="3"/></svg>
            Ver plan
          </a>
          <a href="${esc(borrador_url)}" target="_blank" class="btn-editar">
            <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M11 4H4a2 2 0 0 0-2 2v14a2 2 0 0 0 2 2h14a2 2 0 0 0 2-2v-7"/><path d="M18.5 2.5a2.121 2.121 0 0 1 3 3L12 15l-4 1 1-4 9.5-9.5z"/></svg>
            Editar borrador
          </a>
        </div>
      </div>
    </div>`;
  }).join('');
}

function esc(s) {
  return String(s || '').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;');
}

function mostrarToast(msg, tipo='success') {
  const t = document.getElementById('toast');
  t.textContent = (tipo === 'success' ? '✓ ' : '✕ ') + msg;
  t.className = `toast ${tipo} show`;
  setTimeout(() => t.classList.remove('show'), 3000);
}


function generarCalendario() {
  var desde = document.getElementById('cal-desde').value;
  var hasta = document.getElementById('cal-hasta').value;
  if (!desde || !hasta) {
    mostrarToast('Selecciona mes de inicio y mes final', 'error');
    return;
  }
  var d0 = new Date(desde + '-01');
  var d1 = new Date(hasta + '-01');
  if (d0 > d1) {
    mostrarToast('El mes de inicio debe ser anterior al mes final', 'error');
    return;
  }

  var meses = ['Enero','Febrero','Marzo','Abril','Mayo','Junio',
               'Julio','Agosto','Septiembre','Octubre','Noviembre','Diciembre'];
  var diasSem = ['LU','MA','MI','JU','VI','SA','DO'];
  var colores = [
    {bg:'#c8e6c9',border:'#a5d6a7',label:'Nutrici\u00f3n'},
    {bg:'#bbdefb',border:'#90caf9',label:'Ejercicio'},
    {bg:'#f8bbd0',border:'#f48fb1',label:'Rutina facial'},
    {bg:'#fff9c4',border:'#ffe082',label:'Sue\u00f1o'}
  ];

  var mesesHtml = '';
  var cur = new Date(d0);
  while (cur <= d1) {
    var yr  = cur.getFullYear();
    var mo  = cur.getMonth();
    var diasEnMes = new Date(yr, mo + 1, 0).getDate();
    var primerDia = new Date(yr, mo, 1).getDay();
    primerDia = (primerDia === 0) ? 6 : primerDia - 1;

    var cabHtml = '';
    for (var di = 0; di < diasSem.length; di++) {
      cabHtml += '<div class="cal-dh">' + diasSem[di] + '</div>';
    }

    var celdasHtml = '';
    for (var e = 0; e < primerDia; e++) celdasHtml += '<div class="cal-day empty"></div>';

    for (var dia = 1; dia <= diasEnMes; dia++) {
      var dots = '';
      for (var ci = 0; ci < colores.length; ci++) {
        dots += '<div style="width:11px;height:11px;border-radius:2px;background:' +
                colores[ci].bg + ';border:1px solid ' + colores[ci].border +
                ';flex-shrink:0"></div>';
      }
      celdasHtml += '<div class="cal-day"><div class="cal-day-num">' + dia +
                    '</div><div class="cal-dots">' + dots + '</div></div>';
    }

    var total = primerDia + diasEnMes;
    var rest  = (7 - total % 7) % 7;
    for (var r = 0; r < rest; r++) celdasHtml += '<div class="cal-day empty"></div>';

    mesesHtml += '<div class="cal-month">' +
      '<div class="cal-mhdr">' + meses[mo] + ' ' + yr + '</div>' +
      '<div class="cal-days-hdr">' + cabHtml + '</div>' +
      '<div class="cal-days-grid">' + celdasHtml + '</div>' +
    '</div>';

    cur.setMonth(cur.getMonth() + 1);
  }

  var leyenda = '';
  for (var li = 0; li < colores.length; li++) {
    leyenda += '<div style="display:flex;align-items:center;gap:5px;font-size:11px;color:#6b7280">' +
               '<div style="width:12px;height:12px;border-radius:2px;background:' + colores[li].bg +
               ';border:1px solid ' + colores[li].border + '"></div>' + colores[li].label + '</div>';
  }

  var printHtml =
    '<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8">' +
    '<title>Calendario de Seguimiento \u2014 Centro Carvajal</title>' +
    '<style>' +
    '@page{size:A4 landscape;margin:10mm}' +
    '*{margin:0;padding:0;box-sizing:border-box}' +
    'body{font-family:"Segoe UI",system-ui,sans-serif;background:#faf9f6;color:#1c1c1c}' +
    '.page-header{text-align:center;margin-bottom:14px}' +
    '.page-header h1{font-family:Georgia,serif;font-size:18pt;color:#1c1c1c;margin-bottom:2px}' +
    '.page-header p{font-size:9pt;color:#6b7280}' +
    '.cal-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:10px}' +
    '.cal-month{border:1px solid rgba(143,168,50,.22);border-radius:8px;overflow:hidden;background:#fff}' +
    '.cal-mhdr{background:#1c1c1c;color:#8fa832;text-align:center;padding:7px 10px;font-family:Georgia,serif;font-size:11pt;font-weight:700}' +
    '.cal-days-hdr{display:grid;grid-template-columns:repeat(7,1fr);background:rgba(143,168,50,.1)}' +
    '.cal-dh{text-align:center;font-size:7pt;font-weight:700;color:#8fa832;padding:3px 1px;text-transform:uppercase}' +
    '.cal-days-grid{display:grid;grid-template-columns:repeat(7,1fr)}' +
    '.cal-day{border-right:1px solid rgba(143,168,50,.1);border-bottom:1px solid rgba(143,168,50,.1);padding:3px 4px;background:#fff;min-height:34px}' +
    '.cal-day:nth-child(7n){border-right:none}' +
    '.cal-day.empty{background:#fafafa}' +
    '.cal-day-num{font-size:7.5pt;font-weight:700;color:#1c1c1c;margin-bottom:2px}' +
    '.cal-dots{display:flex;gap:2px;flex-wrap:wrap}' +
    '.leyenda{display:flex;gap:16px;justify-content:center;margin-top:12px;padding-top:8px;border-top:1px solid rgba(143,168,50,.2)}' +
    '.footer{text-align:center;margin-top:10px;font-size:8pt;color:#9aaa8a}' +
    '@media print{body{background:white}}' +
    '</style></head><body>' +
    '<div class="page-header">' +
    '<h1>Calendario de Seguimiento</h1>' +
    '<p>Centro Carvajal \u00b7 M\u00e9todo de Rejuvenecimiento Carvajal \u00b7 Marca cada d\u00eda al completar tu rutina</p>' +
    '</div>' +
    '<div class="cal-grid">' + mesesHtml + '</div>' +
    '<div class="leyenda">' + leyenda + '</div>' +
    '<div class="footer">Centro Carvajal \u00b7 L\u00edderes en Medicina Est\u00e9tica en Panam\u00e1 \u00b7 centrocarvajal.com</div>' +
    '<script>window.onload=function(){window.print();}</scr' + 'ipt>' +
    '</body></html>';

  var win = window.open('', '_blank');
  win.document.write(printHtml);
  win.document.close();
}

// Cargar planes al iniciar
cargarPlanes();
</script>
</body>
</html>""" 


@app.route('/api/login', methods=['POST'])
def api_login():
    data = request.get_json(silent=True) or {}
    pwd    = data.get('password', '').strip()
    stored = ADMIN_PASSWORD.strip()
    if pwd == stored:
        tok = secrets.token_hex(32)
        admin_tokens.add(tok)
        return jsonify({'ok': True, 'token': tok})
    return jsonify({'ok': False}), 401


def _check_token():
    tok = request.headers.get('X-Token', '')
    return tok in admin_tokens


@app.route('/api/planes', methods=['GET'])
def api_listar_planes():
    planes = []
    try:
        import cloudinary.api
        from datetime import datetime as dt
        resultado = cloudinary.api.resources(
            type='upload',
            resource_type='raw',
            prefix='carvajal/planes/',
            max_results=200,
            context=True,
        )
        for r in resultado.get('resources', []):
            public_id      = r.get('public_id', '')
            # secure_url ya incluye la extensión correcta
            plan_url       = r.get('secure_url', '')
            nombre_archivo = public_id.split('/')[-1] + '.html'

            # Extraer campos del nombre de archivo
            # Formato: Plan_Nombre_Apellido_YYYYMMDD_modelo(.html)
            base  = public_id.split('/')[-1]
            base  = re.sub(r'\.html$', '', base, flags=re.IGNORECASE)  # quitar .html si viene
            parts = base.split('_')
            modelo = 'claude'
            fecha_raw = ''
            nombre_parts = []
            for p in parts:
                p_clean = re.sub(r'\.html$', '', p, flags=re.IGNORECASE)  # limpiar .html en cada parte
                if p_clean.lower() in ('claude', 'gemini', 'groq'):
                    modelo = p_clean.lower()
                elif len(p_clean) == 8 and p_clean.isdigit():
                    fecha_raw = p_clean
                elif p_clean not in ('Plan',) and p_clean:
                    nombre_parts.append(p_clean)
            nombre_paciente = ' '.join(nombre_parts) if nombre_parts else base

            # Fecha legible desde nombre de archivo
            fecha_legible = ''
            if fecha_raw and len(fecha_raw) == 8:
                try:
                    fecha_legible = dt.strptime(fecha_raw, '%Y%m%d').strftime('%d/%m/%Y')
                except:
                    fecha_legible = fecha_raw

            # Fecha ISO de Cloudinary para stats "este mes"
            created_at = r.get('created_at', '')

            # job_id: primero desde context (si fue guardado), luego fallback al nombre
            ctx      = r.get('context', {}).get('custom', {})
            job_id_r = ctx.get('job_id', '') or base

            planes.append({
                'url'           : plan_url,
                'nombre_archivo': nombre_archivo,
                'nombre'        : nombre_paciente,
                'modelo'        : modelo,
                'fecha'         : fecha_legible or 'Sin fecha',
                'fecha_raw'     : fecha_raw,
                'job_id'        : job_id_r,
                'created_at'    : created_at,
            })
        planes.sort(key=lambda x: x.get('created_at', ''), reverse=True)
    except Exception as e:
        print(f'[api/planes] Error: {e}')
        return jsonify({'planes': [], 'error': str(e)})
    return jsonify({'planes': planes, 'total': len(planes)})



def generar_analisis_medico(data):
    """Llama a Claude para generar un análisis clínico técnico orientado al médico."""
    print("[analisis_medico] Generando análisis clínico con Claude...")
    t0 = time.time()

    # Construir resumen de datos clínicos relevantes
    campos = {
        'Nombre': data.get('nombre',''),
        'Edad': data.get('edad',''),
        'Sexo': data.get('sexo',''),
        'Fecha nacimiento': data.get('fechaNacimiento',''),
        'IMC': data.get('imc',''),
        'Estatura': data.get('estatura','') + ' cm',
        'Peso': data.get('peso','') + ' kg',
        'Condicion sistemica': data.get('condicionSistemica',''),
        'Otras condiciones': data.get('condiciones',''),
        'Medicamentos': data.get('medicamentos',''),
        'Cirugias': data.get('cirugias',''),
        'Fuma': data.get('fuma',''),
        'Alcohol': data.get('alcohol',''),
        'Embarazo': data.get('embarazo',''),
        'Lactancia': data.get('lactancia',''),
        'Anticonceptivos': data.get('anticonceptivos',''),
        'SOP': data.get('sop',''),
        'Menopausia': data.get('menopausia',''),
        'Perimenopausia': data.get('perimenopausia',''),
        'Alergia lidocaina': data.get('alergia_lidocaina',''),
        'Alergia penicilina': data.get('alergia_penicilina',''),
        'Alergia yodo': data.get('alergia_yodo',''),
        'Alergia AINEs': data.get('alergia_aines',''),
        'Alergia latex': data.get('alergia_latex',''),
        'Alergia aloe': data.get('alergia_aloe',''),
        'Alergia fragancias': data.get('alergia_fragancias',''),
        'Sueno': data.get('sueno',''),
        'Hora duerme': data.get('horaDuerme',''),
        'Hora despierta': data.get('horaDespierta',''),
        'Cansancio diurno': data.get('cansancioDia',''),
        'Nivel estres': data.get('nivelEstres',''),
        'Actividad fisica': data.get('actFisica',''),
        'Evacuacion': data.get('evacuacion',''),
        'Intolerancias': data.get('intolerancias',''),
        'Sintomas digestivos': data.get('sintomasDigestivos',''),
        'Tipo piel': data.get('pielTipo',''),
        'Problemas piel': data.get('pielProblemas',''),
        'Exposicion solar': data.get('solar',''),
        'SPF': data.get('spf',''),
        'Historial estetico': data.get('historialEstetico',''),
        'Laser activo': data.get('laserActivo',''),
        'Contraindicaciones': data.get('contraindications',''),
        'Antecedentes familiares': data.get('antecedentesFam','') + ' ' + data.get('antecedentesFamDet',''),
        'Satisfaccion actual': str(data.get('satisfaccion','')) + '/10',
        'Prioridad': data.get('prioridad',''),
        'Areas faciales': data.get('areasFaciales',''),
        'Areas corporales': data.get('areasCorporales',''),
    }
    def _str(v):
        if isinstance(v, list):  return ', '.join(str(i) for i in v) if v else ''
        if isinstance(v, dict):  return '; '.join(f'{k}: {val}' for k, val in v.items()) if v else ''
        return str(v) if v else ''

    resumen = '\n'.join(
        f'{k}: {_str(v)}'
        for k, v in campos.items()
        if _str(v).strip() not in ('', '0', 'No registrado', ' ')
    )

    system_prompt = """Eres un médico especialista en medicina estética y bienestar integral.
Recibes los datos del cuestionario de un paciente nuevo de Centro Carvajal, clínica de medicina estética en Panamá.
Tu tarea es generar un ANÁLISIS CLÍNICO TÉCNICO para el expediente médico interno — NO para el paciente.
El análisis debe ser en lenguaje médico técnico, objetivo y clínicamente relevante.

Estructura tu respuesta en estas secciones exactas, usando encabezados con ##:

## RESUMEN DEL PERFIL CLÍNICO
2-3 oraciones que resuman el cuadro general del paciente desde perspectiva clínica.

## HALLAZGOS RELEVANTES
Lista de hallazgos clínicamente significativos (metabólicos, hormonales, dermatológicos, nutricionales, etc.).

## POSIBLES CONDICIONES A EVALUAR
Si alguna combinación de síntomas o factores sugiere una condición subyacente, mencionarla con justificación clínica breve. Ser claro que son hipótesis diagnósticas que requieren evaluación, no diagnósticos definitivos.

## CONTRAINDICACIONES Y PRECAUCIONES PARA TRATAMIENTOS ESTÉTICOS
Lista específica de contraindicaciones absolutas y relativas según el perfil del paciente.

## RECOMENDACIONES PARA EL EQUIPO MÉDICO
Acciones sugeridas antes de iniciar tratamientos (exámenes, consultas con especialistas, ajustes de protocolo, etc.).

Sé conciso pero completo. Usa terminología médica adecuada."""

    try:
        resp = req.post(
            'https://api.anthropic.com/v1/messages',
            headers={
                'Content-Type': 'application/json',
                'x-api-key': CLAUDE_KEY,
                'anthropic-version': '2023-06-01',
            },
            json={
                'model': 'claude-opus-4-6',
                'max_tokens': 2000,
                'system': system_prompt,
                'messages': [{'role': 'user', 'content': f'Datos del paciente:\n{resumen}'}],
            },
            timeout=120
        )
        elapsed = round(time.time() - t0, 1)
        if resp.status_code != 200:
            print(f'[analisis_medico] ERROR {resp.status_code}: {resp.text[:200]}')
            return None
        txt = resp.json()['content'][0]['text'].strip()
        print(f'[analisis_medico] OK — {elapsed}s ({len(txt)} chars)')
        return txt
    except Exception as e:
        import traceback
        print(f'[analisis_medico] Excepcion: {e}')
        print(traceback.format_exc())
        return None


def generar_docx_cuestionario(data, plan_json=None, analisis_medico=None):
    """Genera .docx estructurado con los datos del formulario."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    VERDE  = RGBColor(0x2d,0x3a,0x2e)
    OLIVE  = RGBColor(0x8f,0xa8,0x32)
    GOLD   = RGBColor(0xb8,0x93,0x5a)
    BLANCO = RGBColor(0xff,0xff,0xff)
    GRIS   = RGBColor(0x6b,0x72,0x80)

    def set_cell_bg(cell, color_hex):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'), color_hex); tcPr.append(shd)

    def heading(text, level=1):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text); run.bold = True
        run.font.size = Pt(13 if level==1 else 11)
        run.font.color.rgb = OLIVE if level==1 else VERDE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    def row_tabla(tabla, label, value):
        cells = tabla.add_row().cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].bold = True
        cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(cells[0], '2d3a2e')
        cells[0].paragraphs[0].runs[0].font.color.rgb = BLANCO
        val_str = value if isinstance(value,str) else (', '.join(value) if isinstance(value,list) else str(value or '—'))
        cells[1].text = val_str or '—'
        cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    def tabla2():
        t = doc.add_table(rows=0, cols=2); t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        return t

    # Encabezado
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('CENTRO CARVAJAL'); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=VERDE
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Cuestionario del Paciente — Método de Rejuvenecimiento Carvajal')
    r2.font.size=Pt(10); r2.font.color.rgb=GOLD
    doc.add_paragraph()

    # ── Análisis clínico para el médico ──
    if analisis_medico:
        # Título de sección
        p_title = doc.add_paragraph()
        r_title = p_title.add_run('ANÁLISIS CLÍNICO — USO INTERNO DEL EQUIPO MÉDICO')
        r_title.bold = True
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = BLANCO
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after  = Pt(0)
        # Fondo verde oscuro para el título
        pPr = p_title._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1c1c1c'); pPr.append(shd)
        p_title.paragraph_format.left_indent  = Cm(0.4)
        p_title.paragraph_format.right_indent = Cm(0.4)

        # Renderizar secciones del análisis
        import re as _re
        secciones = _re.split(r'\n## ', analisis_medico)
        for i, seccion in enumerate(secciones):
            if not seccion.strip():
                continue
            # Primera sección puede no tener ##
            if i == 0 and not seccion.startswith('##'):
                lineas = seccion.replace('## ', '').strip().split('\n')
            else:
                lineas = seccion.strip().split('\n')

            if not lineas:
                continue

            # Primera línea es el título de subsección
            titulo_sec = lineas[0].strip().replace('## ', '').replace('# ', '')
            p_sec = doc.add_paragraph()
            p_sec.paragraph_format.space_before = Pt(8)
            p_sec.paragraph_format.space_after  = Pt(2)
            r_sec = p_sec.add_run(titulo_sec)
            r_sec.bold = True
            r_sec.font.size = Pt(10)
            r_sec.font.color.rgb = OLIVE

            # Resto del contenido
            cuerpo = '\n'.join(lineas[1:]).strip()
            if cuerpo:
                for linea in cuerpo.split('\n'):
                    linea = linea.strip()
                    if not linea:
                        continue
                    # Limpiar markdown básico
                    linea = _re.sub(r'\*\*(.+?)\*\*', r'\1', linea)
                    linea = linea.lstrip('- •').strip()
                    p_l = doc.add_paragraph(style='List Bullet') if linea else doc.add_paragraph()
                    if linea:
                        run = p_l.add_run(linea)
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = RGBColor(0x1c, 0x1c, 0x1c)
                    p_l.paragraph_format.space_before = Pt(1)
                    p_l.paragraph_format.space_after  = Pt(1)

        doc.add_paragraph()  # espacio antes de datos personales

    heading('1. DATOS PERSONALES')
    t=tabla2()
    row_tabla(t,'nombre_completo',   data.get('nombre',''))
    row_tabla(t,'cedula_pasaporte',  data.get('cedula',''))
    row_tabla(t,'fecha_nacimiento',  data.get('fechaNacimiento',''))
    row_tabla(t,'edad',              data.get('edad',''))
    row_tabla(t,'sexo',              data.get('sexo',''))
    row_tabla(t,'celular',           data.get('celular',''))
    row_tabla(t,'email',             data.get('email',''))
    row_tabla(t,'direccion',         data.get('direccion',''))
    row_tabla(t,'ocupacion',         data.get('ocupacion',''))
    row_tabla(t,'actividad_laboral', data.get('actLaboral',''))
    row_tabla(t,'horario_laboral',   data.get('horarioLaboral',''))
    row_tabla(t,'como_conocio_clinica', data.get('comoConociste',''))
    row_tabla(t,'contacto_emergencia', data.get('contactoEmergencia','') + ' (' + data.get('contactoRelacion','') + ') ' + data.get('contactoTel',''))
    row_tabla(t,'fecha_evaluacion',  data.get('fecha',''))

    heading('2. MEDIDAS CORPORALES')
    t=tabla2()
    row_tabla(t,'estatura_cm',data.get('estatura',''))
    row_tabla(t,'peso_kg',data.get('peso',''))
    row_tabla(t,'imc',data.get('imc',''))

    heading('3. SALUD GENERAL')
    t=tabla2()
    row_tabla(t,'condicion_sistemica',data.get('condicionSistemica',''))
    row_tabla(t,'condiciones_medicas',data.get('condiciones',''))
    row_tabla(t,'medicamentos',data.get('medicamentos',''))
    row_tabla(t,'cirugias_previas',data.get('cirugias',''))
    row_tabla(t,'alergias',data.get('alergias',''))
    row_tabla(t,'fuma',data.get('fuma',''))
    row_tabla(t,'consume_alcohol',data.get('alcohol',''))

    heading('4. EVALUACIÓN CUTÁNEA Y RUTINA FACIAL')
    t=tabla2()
    row_tabla(t,'tipo_piel',data.get('pielTipo',''))
    row_tabla(t,'problemas_faciales',data.get('pielProblemas',''))
    row_tabla(t,'areas_corporales',data.get('areasCorporales',''))
    row_tabla(t,'rutina_manana',data.get('rutinaManana',''))
    row_tabla(t,'rutina_noche',data.get('rutinaNoche',''))
    row_tabla(t,'productos_frecuentes',data.get('productosFrecuentes',''))
    row_tabla(t,'usa_solar',data.get('solar',''))
    row_tabla(t,'spf_factor',data.get('spf',''))
    row_tabla(t,'laser_activo',data.get('laserActivo',''))
    hist=data.get('historialEstetico',[])
    if hist: row_tabla(t,'historial_estetico',hist if isinstance(hist,str) else ', '.join(hist))

    heading('5. ALIMENTACIÓN Y NUTRICIÓN')
    t=tabla2()
    row_tabla(t,'intolerancias_alim',data.get('intolerancias',''))
    row_tabla(t,'sintomas_digestivos',data.get('sintomasDigestivos',''))
    row_tabla(t,'proteinas_frecuentes',data.get('proteinas',''))
    row_tabla(t,'carbohidratos',data.get('carbohidratos',''))
    row_tabla(t,'verduras',data.get('verduras',''))
    row_tabla(t,'frutas',data.get('frutas',''))
    row_tabla(t,'alimentos_evitar',data.get('alimentosEvitar',''))
    row_tabla(t,'postres_dulces',data.get('postres',''))
    row_tabla(t,'bebidas_habituales',data.get('bebidas',''))
    row_tabla(t,'notas_alimentacion',data.get('notasAlimentacion',''))

    heading('6. ACTIVIDAD FÍSICA Y SUEÑO')
    t=tabla2()
    row_tabla(t,'actividad_fisica',data.get('actFisica',''))
    row_tabla(t,'horas_sueno',data.get('sueno',''))
    row_tabla(t,'hora_despierta',data.get('horaDespierta',''))
    row_tabla(t,'hora_duerme',data.get('horaDuerme',''))
    row_tabla(t,'cansancio_diurno',data.get('cansancioDia',''))
    row_tabla(t,'nivel_estres',data.get('nivelEstres',''))
    row_tabla(t,'num_hijos',data.get('numHijos',''))

    heading('7. OBJETIVOS Y PRIORIDADES')
    t=tabla2()
    row_tabla(t,'prioridad_principal',data.get('prioridad',''))
    row_tabla(t,'expectativas',data.get('expectativas',''))
    row_tabla(t,'satisfaccion_actual',str(data.get('satisfaccion',''))+'/10')

    heading('8. ÁREAS DE TRATAMIENTO E HISTORIAL ESTÉTICO')
    t=tabla2()
    row_tabla(t,'areas_faciales',        data.get('areasFaciales',''))
    row_tabla(t,'areas_corporales_detalle', data.get('areasCorporales',''))
    hist    = data.get('historialEstetico', [])
    det_map = data.get('historialDetalle', {})
    if hist:
        for trat in (hist if isinstance(hist, list) else [hist]):
            det = det_map.get(trat, {}) if isinstance(det_map, dict) else {}
            fecha = det.get('fecha', '') if det else ''
            zona  = det.get('zona', '')  if det else ''
            val   = trat
            if fecha: val += f' — última sesión: {fecha}'
            if zona:  val += f' — zona: {zona}'
            row_tabla(t, 'tratamiento', val)
    else:
        row_tabla(t, 'historial_estetico', '—')

    # Contraindicaciones — solo las marcadas como Sí
    contra = data.get('contraindications', {})
    contra_si = [k for k,v in contra.items() if v == 'Si'] if isinstance(contra, dict) else []
    row_tabla(t,'contraindicaciones',    ', '.join(contra_si) if contra_si else 'Ninguna')

    heading('9. ANTECEDENTES FAMILIARES')
    t=tabla2()
    row_tabla(t,'antecedentes_familiares', data.get('antecedentesFam',''))
    row_tabla(t,'detalle_antecedentes',    data.get('antecedentesFamDet',''))

    doc.add_paragraph()
    p_pie = doc.add_paragraph(); p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp = p_pie.add_run('Centro Carvajal · Líderes en Medicina Estética en Panamá · centrocarvajal.com')
    rp.font.size=Pt(8); rp.font.color.rgb=GRIS

    nombre_safe = re.sub(r'[^a-zA-Z0-9]','_', data.get('nombre','Paciente'))
    docx_path = f'/tmp/Cuestionario_{nombre_safe}_{uuid.uuid4().hex[:6]}.docx'
    doc.save(docx_path)
    return docx_path


def worker(job_id, data, faltantes, fotos=None, modelo='claude'):
    try:
        jobs[job_id] = {'status': 'working', 'msg': 'Generando plan con IA (puede tomar 1-2 min)...'}

        plan_json = generar_plan_ia(data, job_id, modelo=modelo)
        if 'error' in plan_json:
            jobs[job_id] = {'status': 'error', 'msg': plan_json['error']}
            return

        jobs[job_id] = {'status': 'working', 'msg': 'Construyendo borrador y plan...'}
        print(f'[worker] Renderizando plan HTML...')
        html = render_plan(plan_json, data)
        print(f'[worker] Plan HTML OK ({len(html)} chars)')
        borrador_html = render_borrador(plan_json, data, job_id)
        print(f'[worker] Borrador HTML OK')

        nombre    = data.get('nombre', 'Paciente')
        if modelo == 'gemini':
            sufijo_modelo = '_gemini'
        elif modelo == 'groq':
            sufijo_modelo = '_groq'
        else:
            sufijo_modelo = '_claude'
        html_name = 'Plan_' + re.sub(r'[^a-zA-Z0-9]', '_', nombre) + '_' + datetime.now().strftime('%Y%m%d') + sufijo_modelo + '.html'
        html_path = os.path.join(PLANES_DIR, html_name)
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html)

        base_url = os.environ.get('BASE_URL', 'https://metodo.centrocarvajal.com')
        html_url_local = f'{base_url}/planes_generados/{html_name}'

        # Subir plan final a Cloudinary
        print(f'[worker] Subiendo plan a Cloudinary...')
        html_url_cdn = subir_plan_cloudinary(html_path, html_name, job_id=job_id)
        print(f'[worker] Cloudinary plan: {html_url_cdn}')
        html_url = html_url_cdn if html_url_cdn else html_url_local

        # Subir borrador editable a Cloudinary
        jobs[job_id] = {'status': 'working', 'msg': 'Guardando borrador editable...'}
        subir_borrador_cloudinary(borrador_html, job_id)
        borrador_url = f'{base_url}/borrador/{job_id}'

        jobs[job_id] = {'status': 'working', 'msg': 'Enviando correos...'}
        fecha_hoy = datetime.now().strftime('%d/%m/%Y a las %H:%M')

        # Generar análisis clínico con Claude (para el médico)
        analisis_medico = None
        try:
            time.sleep(2)  # pausa tras las 3 secciones del plan
            analisis_medico = generar_analisis_medico(data)
        except Exception as e:
            print(f'[worker] Error analisis medico: {e}')

        # Generar .docx del cuestionario para adjuntar
        docx_cuestionario = None
        try:
            docx_cuestionario = generar_docx_cuestionario(data, plan_json=plan_json, analisis_medico=analisis_medico)
            print(f'[worker] Cuestionario .docx: {docx_cuestionario}')
        except Exception as e:
            print(f'[worker] Error docx cuestionario: {e}')

        # Armar lista de adjuntos: fotos + docx
        adjuntos = list(fotos or [])
        if docx_cuestionario and os.path.exists(docx_cuestionario):
            adjuntos.append(docx_cuestionario)

        # Correo al staff con borrador + adjuntos + CC
        enviar_resend(
            f'Nuevo Plan IA - {nombre} ({fecha_hoy})',
            email_formulario(data, faltantes, borrador_url),
            MAIL_TO,
            adjuntos_extra=adjuntos,
            cc=MAIL_CC or None
        )

        # Limpiar temporales
        for f_path in adjuntos:
            try: os.unlink(f_path)
            except: pass

        jobs[job_id] = {
            'status'      : 'done',
            'nombre'      : nombre,
            'html_url'    : html_url,
            'html_name'   : html_name,
            'borrador_url': borrador_url,
            'faltantes'   : faltantes,
        }

    except Exception as e:
        import traceback
        print(f'[worker] EXCEPCION: {e}')
        print(traceback.format_exc())
        jobs[job_id] = {'status': 'error', 'msg': str(e)}


# ════════════════════════════════════════════════════════════
# LEER DOCX
# ════════════════════════════════════════════════════════════

def leer_docx(path):
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read('word/document.xml').decode('utf-8')
        texto = re.sub(r'</w:p>', '\n', xml)
        texto = re.sub(r'</w:tc>', '\t', texto)
        texto = re.sub(r'<[^>]+>', '', texto)
        texto = htmllib.unescape(texto)
        texto = re.sub(r'[ ]+', ' ', texto)
        lineas = [l.strip() for l in texto.split('\n') if l.strip()]
        pares = []
        for i in range(0, len(lineas) - 1, 2):
            key = lineas[i].lower().strip()
            val = lineas[i + 1].strip()
            if re.match(r'^[a-z][a-z0-9_]+$', key) and val:
                pares.append(f'{key}\t{val}')
        return '\n'.join(pares)
    except Exception as e:
        print(f'leer_docx error: {e}')
        return None




# ════════════════════════════════════════════════════════════
# MAPEAR FORMULARIO WEB → data dict
# ════════════════════════════════════════════════════════════

def _mapear_formulario(f):
    """Convierte el JSON de index.html al mismo dict que usan generar_plan_ia() y render_plan()."""

    def s(key, default=''):
        val = f.get(key, default)
        if isinstance(val, list):
            return ', '.join(str(v) for v in val if v)
        return str(val).strip() if val else default

    def lst(key):
        val = f.get(key, [])
        return val if isinstance(val, list) else []

    est = s('estatura')
    pes = s('peso')
    imc = 'No registrado'
    if est and pes:
        try: imc = f'{float(pes) / (float(est)/100)**2:.1f}'
        except: pass

    contra_raw = f.get('contraindications', {})
    contra = {k: ('Si' if str(v).lower() in ['si','sí','yes'] else 'No') for k, v in contra_raw.items()}

    intolerancias = lst('intolerancias')
    sintomas_map = {
        'hinchazon_abdominal': 'Hinchazon',
        'gases':               'Gases',
        'estrenimiento':       'Estrenimiento',
        'cansancio_comidas':   'Cansancio tras comer',
        'digestion_lenta':     'Digestion lenta',
        'nauseas':             'Nauseas',
    }
    sintomas = [sintomas_map[k] for k in lst('sintomasDigestivos') if k in sintomas_map]
    faciales   = lst('areasFaciales')
    corporales = lst('areasCorporales')
    rutina_m = s('rutinaManana')
    rutina_n = s('rutinaNoche')
    rutina_facial = (rutina_m + ' | ' + rutina_n).strip(' |') if (rutina_m or rutina_n) else 'No tiene rutina facial'

    return {
        'nombre':              s('nombre'),
        'cedula':              s('cedula'),
        'direccion':           s('direccion'),
        'edad':                s('edad'),
        'fechaNacimiento':     s('fechaNacimiento'),
        'sexo':                s('sexo'),
        'ocupacion':           s('ocupacion'),
        'actLaboral':          s('actLaboral'),
        'horarioLaboral':      s('horarioLaboral'),
        'email':               s('email'),
        'celular':             s('celular'),
        'comoConociste':       s('comoConociste'),
        'fecha':               datetime.now().strftime('%d de %B, %Y'),
        'estatura':            est or None,
        'peso':                pes or None,
        'imc':                 imc,
        'pielTipo':            s('pielTipo'),
        'pielProblemas':       faciales,
        'rutinaFacial':        rutina_facial,
        'rutinaManana':        rutina_m,
        'rutinaNoche':         rutina_n,
        'productosFrecuentes': s('productosFrecuentes'),
        'solar':               s('solar'),
        'spf':                 s('spf'),
        'actFisica':           s('actFisica'),
        'sueno':               s('sueno'),
        'horaDespierta':       s('horaDespierta'),
        'horaDuerme':          s('horaDuerme'),
        'cansancioDia':        'Si' if s('cansancioDia').lower() in ['si','sí','yes'] else 'No',
        'fuma':                s('fuma'),
        'alcohol':             s('alcohol'),
        'condicionSistemica':  s('condicionSistemica') or 'Sin enfermedades',
        'condiciones':         s('condiciones'),
        'medicamentos':        s('medicamentos'),
        'cirugias':            s('cirugias') or 'Ninguna',
        'antecedentesFam':     'Ninguno',
        'alergias':            s('alergias') or 'Ninguna',
        'alergiasDetalle':     f.get('alergiasDetalle', {}),
        'contraindications':   contra,
        'embarazo':            s('embarazo'),
        'lactancia':           s('lactancia'),
        'anticonceptivos':     s('anticonceptivos'),
        'sop':                 s('sop'),
        'menopausia':          s('menopausia'),
        'perimenopausia':      s('perimenopausia'),
        'alergia_lidocaina':   s('alergia_lidocaina'),
        'alergia_penicilina':  s('alergia_penicilina'),
        'alergia_yodo':        s('alergia_yodo'),
        'alergia_aines':       s('alergia_aines'),
        'alergia_latex':       s('alergia_latex'),
        'alergia_aloe':        s('alergia_aloe'),
        'alergia_fragancias':  s('alergia_fragancias'),
        'evacuacion':          s('evacuacion'),
        'antecedentesFam':     s('antecedentesFam'),
        'antecedentesFamDet':  s('antecedentesFamDet'),
        'comoConociste':       s('comoConociste'),
        'areasFaciales':       faciales,
        'areasCorporales':     corporales,
        'prioridad':           s('prioridad'),
        'expectativas':        s('expectativas'),
        'satisfaccion':        s('satisfaccion'),
        'historialEstetico':   lst('historialEstetico'),
        'historialDetalle':    f.get('historialDetalle', {}),
        'laserActivo':         'Si' if s('laserActivo').lower() in ['si','sí','yes'] else 'No',
        'intolerancias':       intolerancias,
        'sintomasDigestivos':  sintomas,
        'proteinas':           s('proteinas'),
        'carbohidratos':       s('carbohidratos'),
        'verduras':            s('verduras'),
        'frutas':              s('frutas'),
        'alimentosEvitar':     s('alimentosEvitar'),
        'postres':             s('postres'),
        'bebidas':             s('bebidas'),
        'notasAlimentacion':   s('notasAlimentacion'),
        'nivelEstres':         s('nivelEstres'),
        'numHijos':            s('numHijosVal') or s('numHijos'),
        'notasStaff':          '',
    }

# ════════════════════════════════════════════════════════════
# PARSEAR CUESTIONARIO
# ════════════════════════════════════════════════════════════

def parsear_cuestionario(texto):
    raw = {}
    for linea in texto.split('\n'):
        linea = linea.strip()
        if not linea: continue
        if '\t' in linea:
            parts = linea.split('\t', 1)
            key, val = parts[0], parts[1]
        else:
            m = re.match(r'^([a-z][a-z0-9_]+)\s+(.+)$', linea, re.I)
            if m: key, val = m.group(1), m.group(2)
            else: continue
        key, val = key.lower().strip(), val.strip()
        if key and val: raw[key] = val

    def v(k):
        val = raw.get(k)
        if not val or val == '0' or val.lower() in ['ninguna','ninguno','aun no','n/a','nada']:
            return None
        return val

    def si(k):
        return raw.get(k, 'NO').upper() in ['SI', 'SÍ', 'S', 'YES']

    faltantes = []
    for campo, label in [('peso','Peso'),('altura','Altura'),('piel_tipo','Tipo de piel'),
                          ('prioridad_principal','Prioridad principal'),('satisfaccion','Satisfaccion (1-10)'),
                          ('expectativas','Expectativas'),('exposicion_solar','Exposicion solar')]:
        if not raw.get(campo) or raw.get(campo) == '0':
            faltantes.append(label)

    historial = []
    for k, l in [('botox','Botox'),('rellenos','Rellenos'),('hilos','Hilos PDO'),('peeling','Peeling'),
                 ('laser','Laser'),('microblading','Microblading'),('radiofrecuencia','Radiofrecuencia')]:
        if si('tratamiento_' + k): historial.append(l)

    contra = {}
    for c, l in [('embarazada','Embarazo'),('lactancia','Lactancia'),('anticonceptivos','Anticonceptivos'),
                 ('sop','SOP'),('menopausia','Menopausia'),('fuma','Tabaquismo'),
                 ('alergia_lidocaina','Alergia lidocaina'),('alergia_penicilina','Alergia penicilina'),
                 ('alergia_yodo','Alergia yodo'),('alergia_aines','Alergia AINEs'),
                 ('alergia_latex','Alergia latex')]:
        contra[l] = 'Si' if si(c) else 'No'

    sintomas = []
    for c, l in [('hinchazon_abdominal','Hinchazon'),('gases_flatulencias','Gases'),
                 ('estrenimiento_diarrea','Estrenimiento'),('cansancio_comidas','Cansancio tras comer'),
                 ('digestion_lenta','Digestion lenta'),('nauseas_malestar','Nauseas')]:
        if si(c): sintomas.append(l)

    intol = []
    if si('sintomas_lacteos'):    intol.append('Lacteos')
    if si('sintomas_gluten'):     intol.append('Gluten')
    if si('sintomas_procesados'): intol.append('Procesados')

    faciales, corporales = [], []
    if raw.get('area_flacidez_facial'): faciales.append('Flacidez facial')
    if raw.get('area_grasa'):           corporales.append('Grasa localizada')
    if raw.get('area_arrugas'):         faciales.append('Arrugas')
    if raw.get('area_manchas'):         faciales.append('Manchas')
    if raw.get('area_ojeras'):          faciales.append('Ojeras')
    if raw.get('area_celulitis'):       corporales.append('Celulitis')
    prio = (v('prioridad_principal') or '').lower()
    if 'rostro' in prio or 'cara' in prio: faciales.append('Rejuvenecimiento facial')
    if 'cabello' in prio or 'pelo' in prio: faciales.append('Caida de cabello')
    if 'mancha' in prio or 'melasma' in prio: faciales.append('Manchas')
    faciales   = list(dict.fromkeys(faciales))
    corporales = list(dict.fromkeys(corporales))

    ale_list = []
    for c, l in [('alergia_lidocaina','Lidocaina'),('alergia_penicilina','Penicilina'),
                 ('alergia_yodo','Yodo'),('alergia_aines','AINEs'),('alergia_latex','Latex'),
                 ('alergia_aloe','Aloe'),('alergia_fragancias','Fragancias')]:
        if si(c): ale_list.append(l)
    if si('alergia_medicamentos'): ale_list.append(v('medicamentos_cuales') or 'Medicamentos')

    est = v('altura') or v('talla') or v('estatura')
    pes = v('peso')
    imc = 'No registrado'
    if est and pes:
        try: imc = f'{float(pes) / (float(est)/100)**2:.1f}'
        except: pass

    data = {
        'nombre':              v('nombre_completo') or '',
        'edad':                v('edad') or '',
        'sexo':                v('sexo') or '',
        'ocupacion':           v('ocupacion') or v('profesion_trabajo') or '',
        'actLaboral':          v('nivel_actividad_laboral') or '',
        'horarioLaboral':      v('horario_laboral') or '',
        'email':               v('email') or '',
        'fecha':               datetime.now().strftime('%d de %B, %Y'),
        'estatura':            est,
        'peso':                pes,
        'imc':                 imc,
        'pielTipo':            v('piel_tipo') or v('tipo_piel') or '',
        'pielProblemas':       faciales,
        'rutinaFacial':        (v('rutina_manana') or '') + ' | ' + (v('rutina_noche') or '') if si('rutina_diaria_cuidado') else 'No tiene rutina facial',
        'rutinaManana':        v('rutina_manana') or '',
        'rutinaNoche':         v('rutina_noche') or '',
        'productosFrecuentes': v('productos_cosmeticos_frecuentes') or '',
        'solar':               v('exposicion_solar') or '',
        'spf':                 (v('protector_fps') or '') + ' ' + (v('protector_marca') or ''),
        'actFisica':           v('actividad_fisica') or '',
        'sueno':               (v('horas_sueno') or '') + ' horas ' + (v('calidad_sueno') or ''),
        'horaDespierta':       v('hora_levanta') or '',
        'horaDuerme':          v('hora_acuesta') or '',
        'cansancioDia':        'Si' if si('cansancio_dia') else 'No',
        'fuma':                ('Si - ' + (v('fuma_cantidad') or '')) if si('fuma') else 'No',
        'alcohol':             'Si' if si('alcohol') else 'No',
        'condicionSistemica':  (v('enfermedad_detalle') or 'Si') if si('sufre_enfermedad') else 'Sin enfermedades',
        'condiciones':         v('otras_condiciones') or '',
        'medicamentos':        v('medicamento_cual') or '',
        'cirugias':            (v('cirugias_detalle') or 'Si') if si('cirugias') else 'Ninguna',
        'antecedentesFam':     (v('antecedentes_detalle') or 'Si') if si('antecedentes_familiares') else 'Ninguno',
        'alergias':            ', '.join(ale_list) if ale_list else 'Ninguna',
        'contraindications':   contra,
        'areasFaciales':       faciales,
        'areasCorporales':     corporales,
        'prioridad':           v('prioridad_principal') or '',
        'expectativas':        v('expectativas') or '',
        'satisfaccion':        v('satisfaccion') or '',
        'historialEstetico':   list(dict.fromkeys(historial)),
        'laserActivo':         'Si' if si('laser_actual') else 'No',
        'intolerancias':       intol,
        'sintomasDigestivos':  sintomas,
        'proteinas':           ', '.join(filter(None, [v('proteina_pollo'), v('proteinas_otras')])),
        'carbohidratos':       v('carb_arroz_blanco') or '',
        'verduras':            v('verduras_consume') or '',
        'frutas':              v('carb_frutas') or '',
        'alimentosEvitar':     ', '.join(filter(None, [v('proteinas_evitar'), v('carbohidratos_evitar'), v('verduras_evitar')])),
        'postres':             v('postres_favoritos') or '',
        'bebidas':             (v('bebidas_azucaradas_cuales') or 'Si') if si('bebidas_azucaradas') else 'No',
        'notasAlimentacion':   v('observaciones_alimentarias') or '',
        'notasStaff':          '',
    }
    return data, faltantes


# ════════════════════════════════════════════════════════════
# GENERAR PLAN CON CLAUDE API — 3 LLAMADAS SEGMENTADAS
# ════════════════════════════════════════════════════════════

import time

def _datos_paciente(d):
    est = d.get('estatura', '')
    pes = d.get('peso', '')
    imc = d.get('imc', 'No registrado')
    contra_activas = [k for k, v in d.get('contraindications', {}).items() if v == 'Si']
    contra_txt = 'CONTRAINDICACIONES ACTIVAS: ' + ', '.join(contra_activas) if contra_activas else 'Sin contraindicaciones activas.'
    return f"""DATOS DEL PACIENTE:
Nombre: {d['nombre']} | Edad: {d['edad']} | Sexo: {d['sexo']}
Ocupacion: {d['ocupacion']} | Horario: {d['horarioLaboral']} | Act.laboral: {d['actLaboral']}
Fecha evaluacion: {d['fecha']}
Estatura: {est}cm | Peso: {pes}kg | IMC: {imc}

PIEL: {d['pielTipo']} | Problemas: {', '.join(d['pielProblemas'])}
Rutina manana: {d['rutinaManana']} | Noche: {d['rutinaNoche']}
Productos: {d['productosFrecuentes']} | Solar: {d['solar']} | SPF: {d['spf']}

HABITOS: Act.fisica: {d['actFisica']} | Sueno: {d['sueno']}
Fuma: {d['fuma']} | Alcohol: {d['alcohol']}

SALUD: {contra_txt}
Condicion sistemica: {d['condicionSistemica']}
Condiciones: {d['condiciones']}
Medicamentos: {d['medicamentos']}
Cirugias: {d['cirugias']}
Alergias: {d['alergias']}

OBJETIVOS: Faciales: {', '.join(d['areasFaciales'])} | Corporales: {', '.join(d['areasCorporales'])}
Prioridad (palabras del paciente): "{d['prioridad']}"
Expectativas: {d['expectativas']} | Satisfaccion: {d['satisfaccion']}/10
Historial estetico: {', '.join(d['historialEstetico']) or 'Ninguno'}

ALIMENTACION:
Intolerancias: {', '.join(d['intolerancias']) or 'Ninguna'}
Proteinas: {d['proteinas']} | Carbos: {d['carbohidratos']}
Verduras: {d['verduras']} | Frutas: {d['frutas']}
Evitar: {d['alimentosEvitar']}
Notas: {d['notasAlimentacion']}

CONTEXTO PERSONAL ADICIONAL:
Numero de hijos: {d.get('numHijos','No especificado')}
Nivel de estres (1-10): {d.get('nivelEstres','No especificado')}"""


def _llamar_claude(num, total, system_prompt, user_msg, max_tok=6000):
    print(f"[{num}/{total}] Iniciando...")
    t0 = time.time()
    resp = req.post(
        'https://api.anthropic.com/v1/messages',
        headers={
            'Content-Type': 'application/json',
            'x-api-key': CLAUDE_KEY,
            'anthropic-version': '2023-06-01',
        },
        json={
            'model': 'claude-sonnet-4-6',
            'max_tokens': max_tok,
            'system': system_prompt,
            'messages': [{'role': 'user', 'content': user_msg}],
        },
        timeout=300
    )
    elapsed = round(time.time() - t0, 1)

    if resp.status_code != 200:
        print(f"[{num}/{total}] ERROR {resp.status_code}: {resp.text[:400]}")
        return None, f'Error API Claude ({resp.status_code}): {resp.text[:300]}'

    rj = resp.json()
    txt = rj['content'][0]['text']
    stop_reason = rj.get('stop_reason', '')
    usage = rj.get('usage', {})
    tok_in  = usage.get('input_tokens', '?')
    tok_out = usage.get('output_tokens', '?')
    print(f"[{num}/{total}] OK — {elapsed}s | input: {tok_in} tokens | output: {tok_out} tokens | stop: {stop_reason}")

    if stop_reason == 'max_tokens':
        return None, f'Llamada {num} truncada por limite de tokens.'

    txt = re.sub(r'^```json\s*', '', txt.strip())
    txt = re.sub(r'^```\s*', '', txt)
    txt = re.sub(r'```\s*$', '', txt).strip()

    try:
        result = json.loads(txt)
        print(f"[{num}/{total}] JSON OK — claves: {list(result.keys())}")
        return result, None
    except Exception as e:
        print(f"[{num}/{total}] JSON invalido: {e} | inicio: {txt[:300]}")
        return None, f'JSON invalido en llamada {num}: {str(e)[:150]}'



def _llamar_gemini(num, total, system_prompt, user_msg, max_tok=8000):
    """Llama a Gemini 2.0 Flash con el mismo contrato que _llamar_claude."""
    print(f"[Gemini {num}/{total}] Iniciando...")
    t0 = time.time()
    model = 'gemini-2.5-flash-lite'
    url = f'https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={GEMINI_KEY}'
    payload = {
        'system_instruction': {'parts': [{'text': system_prompt}]},
        'contents': [{'role': 'user', 'parts': [{'text': user_msg}]}],
        'generationConfig': {
            'maxOutputTokens': max_tok,
            'temperature': 0.7,
        }
    }
    try:
        resp = req.post(url, json=payload, timeout=300)
        elapsed = round(time.time() - t0, 1)
        if resp.status_code != 200:
            print(f"[Gemini {num}/{total}] ERROR {resp.status_code}: {resp.text[:300]}")
            return None, f'Error API Gemini ({resp.status_code}): {resp.text[:200]}'
        rj = resp.json()
        txt = rj['candidates'][0]['content']['parts'][0]['text']
        print(f"[Gemini {num}/{total}] OK — {elapsed}s")
        txt = re.sub(r'^```json\s*', '', txt.strip())
        txt = re.sub(r'^```\s*', '', txt)
        txt = re.sub(r'```\s*$', '', txt).strip()
        try:
            result = json.loads(txt)
            print(f"[Gemini {num}/{total}] JSON OK — claves: {list(result.keys())}")
            return result, None
        except Exception as e:
            print(f"[Gemini {num}/{total}] JSON invalido: {e} | inicio: {txt[:300]}")
            return None, f'JSON invalido en llamada Gemini {num}: {str(e)[:150]}'
    except Exception as e:
        return None, f'Error llamando Gemini: {str(e)[:200]}'



def _llamar_groq(num, total, system_prompt, user_msg, max_tok=8000):
    """Llama a Groq (Llama 3.3 70B) con el mismo contrato que _llamar_claude."""
    print(f"[Groq {num}/{total}] Iniciando...")
    t0 = time.time()
    url = 'https://api.groq.com/openai/v1/chat/completions'

    # Prefijo que fuerza respuestas detalladas y completas
    prefijo = (
        "INSTRUCCION CRITICA DE CALIDAD: Eres un experto medico y de salud integral. "
        "Debes ser EXTREMADAMENTE detallado, especifico y personalizado. "
        "PROHIBIDO usar frases genericas o vagas. "
        "Cada campo del JSON debe tener minimo 2-3 oraciones ricas en contenido clinico. "
        "El menu semanal debe tener comidas COMPLETAS y VARIADAS para cada dia. "
        "Los protocolos deben tener pasos numerados con tiempos y cantidades exactas. "
        "Las recomendaciones deben mencionar el nombre del paciente y sus condiciones especificas. "
        "Completa ABSOLUTAMENTE TODOS los campos sin omitir ninguno. "
        "La calidad de tu respuesta es lo mas importante — tomaté el tiempo necesario.\n\n"
    )

    payload = {
        'model': 'llama-3.3-70b-versatile',
        'messages': [
            {'role': 'system', 'content': prefijo + system_prompt},
            {'role': 'user',   'content': user_msg}
        ],
        'max_tokens': max_tok,
        'temperature': 0.4,  # más bajo = más preciso y consistente
    }
    try:
        resp = req.post(
            url,
            headers={'Authorization': f'Bearer {GROQ_KEY}', 'Content-Type': 'application/json'},
            json=payload,
            timeout=300
        )
        elapsed = round(time.time() - t0, 1)
        if resp.status_code != 200:
            print(f"[Groq {num}/{total}] ERROR {resp.status_code}: {resp.text[:300]}")
            return None, f'Error API Groq ({resp.status_code}): {resp.text[:200]}'
        rj = resp.json()
        txt = rj['choices'][0]['message']['content']
        print(f"[Groq {num}/{total}] OK — {elapsed}s")
        txt = re.sub(r'^```json\s*', '', txt.strip())
        txt = re.sub(r'^```\s*', '', txt)
        txt = re.sub(r'```\s*$', '', txt).strip()
        try:
            result = json.loads(txt)
            print(f"[Groq {num}/{total}] JSON OK — claves: {list(result.keys())}")
            return result, None
        except Exception as e:
            print(f"[Groq {num}/{total}] JSON invalido: {e} | inicio: {txt[:300]}")
            return None, f'JSON invalido en llamada Groq {num}: {str(e)[:150]}'
    except Exception as e:
        return None, f'Error llamando Groq: {str(e)[:200]}'






def generar_plan_ia(d, job_id=None, modelo='claude'):
    datos = _datos_paciente(d)
    t_total = time.time()

    def actualizar(msg, pct=None):
        if job_id and job_id in jobs:
            jobs[job_id]['msg'] = msg
            if pct is not None:
                jobs[job_id]['pct'] = pct

    SYS1 = '''Eres el generador de contenido para planes del METODO CARVAJAL.
Devuelve UNICAMENTE JSON valido sin explicaciones ni markdown.
Genera SOLO estas 3 claves: portada, diagnostico, rutina.
{"portada":{"titulo_pilares":"Los 5 Pilares - [condicion]","intro":"3-4 lineas calido motivador","pilares_resumen":[{"num":1,"emoji":"\ud83e\udd57","titulo":"Titulo adaptado","descripcion":"2-3 lineas"},{"num":2,"emoji":"\ud83c\udfc3","titulo":"Titulo","descripcion":"2-3 lineas"},{"num":3,"emoji":"\ud83e\udde0","titulo":"Titulo","descripcion":"2-3 lineas"},{"num":4,"emoji":"\ud83d\ude34","titulo":"Titulo","descripcion":"2-3 lineas"},{"num":5,"emoji":"\u2728","titulo":"Titulo","descripcion":"2-3 lineas"}]},"diagnostico":{"nota_medica":"Nota alertas criticas","filas":[{"area":"Antropometria","estado":"datos","hallazgos":"analisis","alerta":"normal"},{"area":"Salud Digestiva","estado":"...","hallazgos":"...","alerta":"normal"},{"area":"Sueno y Energia","estado":"...","hallazgos":"...","alerta":"normal"},{"area":"Evaluacion Cutanea","estado":"...","hallazgos":"...","alerta":"normal"},{"area":"Salud Capilar","estado":"...","hallazgos":"...","alerta":"normal"},{"area":"Prioridad Principal","estado":"palabras textuales","hallazgos":"...","alerta":"normal"},{"area":"Condiciones Medicas","estado":"...","hallazgos":"...","alerta":"normal"},{"area":"Rutina Facial Actual","estado":"...","hallazgos":"...","alerta":"normal"},{"area":"Estilo de Vida","estado":"...","hallazgos":"...","alerta":"normal"}]},"rutina":{"nota":"nota rutina","items":[{"hora":"07:00","actividad":"descripcion","pilar":"Nutricion"}]}}
REGLA: Maximo 8 items en rutina. Tips hiperspecificos con nombre y profesion.'''

    SYS2 = '''Eres el generador de contenido para planes del METODO CARVAJAL.
Devuelve UNICAMENTE JSON valido sin explicaciones ni markdown.
Genera SOLO estas 3 claves: pilar1, pilar2, pilar3.
{"pilar1":{"titulo":"Nutricion adaptada","objetivo":"2-3 lineas","frase_motivacional":"frase corta","frase_posicion":"inicio","permitidos":["item"],"evitar":["item"],"menu":[{"dia":"Lunes","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."},{"dia":"Martes","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."},{"dia":"Miercoles","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."},{"dia":"Jueves","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."},{"dia":"Viernes","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."},{"dia":"Sabado","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."},{"dia":"Domingo","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."}],"compras":[{"categoria":"Proteinas","emoji":"\ud83e\udd69","items":["i1","i2","i3","i4","i5"]},{"categoria":"Carbohidratos","emoji":"\ud83c\udf3e","items":["i1","i2","i3","i4"]},{"categoria":"Vegetales","emoji":"\ud83e\udd66","items":["i1","i2","i3","i4","i5"]},{"categoria":"Frutas","emoji":"\ud83c\udf4e","items":["i1","i2","i3","i4"]},{"categoria":"Grasas","emoji":"\ud83e\udd51","items":["i1","i2","i3"]},{"categoria":"Otros","emoji":"\ud83e\uddf4","items":["i1","i2","i3","i4"]}],"suplementacion":["Sup1: dosis"],"tips":[{"texto":"tip especifico con nombre"}]},"pilar2":{"titulo":"Actividad Fisica","objetivo":"objetivo","frase_motivacional":"frase","frase_posicion":"medio","plan_semanal":"plan dia a dia","adaptaciones":"adaptaciones","tips":[{"texto":"tip"}]},"pilar3":{"titulo":"Bienestar Mental","objetivo":"objetivo","frase_motivacional":"frase","frase_posicion":"final","tecnicas":["t1","t2","t3","t4","t5"],"tips":[{"texto":"tip"}]}}
REGLAS: Respetar intolerancias. Tips con nombre, profesion, horario real.'''

    CATALOGO = (
        'BASE DE CONOCIMIENTO CLINICO - TRATAMIENTOS CENTRO CARVAJAL\n'
        'Usa UNICAMENTE estos tratamientos con precios y datos clinicos completos.\n\n'
        '[Ultra Focus (HIFU)] Precios: Fine Lift paquete $999 / Skin Lift Pro 1ses $600 / Total Lift paquete $1300 | Problemas: Flacidez facial,Flacidez corporal,Papada,Redefinición contorno | Zonas: Rostro,Cuello,Papada,Abdomen,Brazos | Grado: Moderado,Severo | Sesiones: 1-2 | Intervalo: 6-12 meses | Recovery: 24-72h leve edema/rojeces | Combinar: RF microagujas,Exilis | NO combinar: Calor profundo intenso mismo día | Orden: Primero HIFU, complementarios en semanas | CONTRAINDICACIONES: Embarazo,Marcapasos,Cáncer activo,Infección activa,Heridas abiertas\n'
        '[Exilis] Precios: EXILIS Abdomen 8ses $2000 | Problemas: Grasa localizada,Flacidez corporal,Contorno corporal,Celulitis | Zonas: Abdomen,Flancos,Brazos,Piernas,Glúteos | Grado: Leve,Moderado | Sesiones: 4 | Intervalo: 10-14 días | Recovery: Sin downtime | Combinar: Presoterapia,X-Wave | NO combinar: Calor profundo intenso mismo día | Orden: Exilis → drenaje (mismo día o posterior) | CONTRAINDICACIONES: Embarazo,Marcapasos,Cáncer activo,Infección activa,Heridas abiertas\n'
        '[Celulite Shock] Precios: Cellulite Shock 10ses $790 | Problemas: Celulitis,Fibrosis,Textura de piel | Zonas: Piernas,Muslos,Glúteos,Caderas | Grado: Leve,Moderado,Severo | Sesiones: 6-10 | Intervalo: 1 semana | Recovery: Sin downtime | Combinar: Presoterapia,Exilis | NO combinar: — | Orden: X-Wave → drenaje / remodelación | CONTRAINDICACIONES: Embarazo,Cáncer activo,Infección activa,Hematomas severos\n'
        '[Acthyderm (Electroporación)] Precios: Acthyderm Rostro 3ses $334 / Acthyderm Cuerpo 12ses $783 | Problemas: Hidratación/Opacidad,Poros/Textura,Antiedad,Manchas | Zonas: Rostro,Cuello,Brazos,Cuerpo | Grado: Leve,Moderado | Sesiones: 4-8 | Intervalo: 1-2 semanas | Recovery: Sin downtime | Combinar: Hidrofacial,Peelings cosméticos | NO combinar: — | Orden: Limpieza → activos → electroporación | CONTRAINDICACIONES: Embarazo,Marcapasos,Epilepsia,Cáncer activo,Infección activa\n'
        '[Radiofrecuencia fraccionada con microagujas] Precios: De Age Treatment paquete $975 | Problemas: Arrugas/Líneas,Poros/Textura,Cicatrices acné,Estrías,Flacidez facial | Zonas: Rostro,Cuello,Abdomen,Glúteos | Grado: Moderado,Severo | Sesiones: 3-5 | Intervalo: 4-6 semanas | Recovery: 3-10 días (según intensidad) | Combinar: Hidrofacial,Acthyderm | NO combinar: Peelings fuertes inmediato | Orden: Preparación → RF → reparación | CONTRAINDICACIONES: Embarazo,Marcapasos,Infección activa,Alteración coagulación,Cáncer activo\n'
        '[Depilación IPL] Precios: IPL Facial 6ses $350 / IPL Axilas $350 / IPL Piernas 8ses $650 / IPL Brasileno 8ses $600 | Problemas: Vello no deseado | Zonas: Rostro,Cuerpo | Grado: Leve,Moderado | Sesiones: 6-10 | Intervalo: 4-6 semanas | Recovery: Sin downtime | Combinar: — | NO combinar: Bronceado reciente | Orden: Evaluación → sesiones periódicas | CONTRAINDICACIONES: Embarazo,Fotosensibilizantes,Infección activa,Piel lesionada,Sol/broncedo reciente\n'
        '[Presoterapia] Precios: Consultar | Problemas: Retención de líquidos/Edema,Postoperatorio,Celulitis,Piernas cansadas | Zonas: Piernas,Abdomen,Brazos | Grado: Leve,Moderado | Sesiones: 6-10 | Intervalo: 1 semana | Recovery: Sin downtime | Combinar: X-Wave,Exilis | NO combinar: — | Orden: Drenaje post tratamiento | CONTRAINDICACIONES: Trombosis/TVP,Insuficiencia arterial,Cáncer activo,Infección activa,Embarazo\n'
        '[Toxina Botulínica] Precios: Toxina Botulinica 30u $450 / 50u $750 | Problemas: Arrugas/Líneas,Rejuvenecimiento facial | Zonas: Rostro | Grado: Leve,Moderado | Sesiones: Según criterio médico | Intervalo: 3-6 meses | Recovery: Mínimo | Combinar: Rellenos,Hilos,Peelings | NO combinar: — | Orden: Evaluación médica → aplicación | CONTRAINDICACIONES: Embarazo,Lactancia,Infección activa,Enfermedad neuromuscular (según criterio médico)\n'
        '[Rellenos Dérmicos] Precios: Plasma Gel 1ses $250 | Problemas: Pérdida de volumen,Arrugas/Líneas,Contorno facial | Zonas: Rostro | Grado: Leve,Moderado,Severo | Sesiones: Según criterio médico | Intervalo: — | Recovery: Mínimo | Combinar: Toxina Botulínica,Hilos | NO combinar: — | Orden: Evaluación médica → plan → aplicación | CONTRAINDICACIONES: Embarazo,Lactancia,Infección activa,Alteración coagulación (según criterio médico)\n'
        '[Hilos PDO] Precios: Hilos PDO 1ses $800 | Problemas: Flacidez facial,Redefinición contorno | Zonas: Rostro,Cuello | Grado: Moderado,Severo | Sesiones: Según criterio médico | Intervalo: — | Recovery: Días (según caso) | Combinar: Toxina Botulínica,Rellenos,HIFU | NO combinar: — | Orden: Plan médico → procedimiento → seguimiento | CONTRAINDICACIONES: Embarazo,Lactancia,Infección activa,Alteración coagulación (según criterio médico)\n'
        '[Plasma / Plasma Gel] Precios: Plasma Facial 1ses $200 / Plasma Gel 1ses $250 | Problemas: Rejuvenecimiento facial,Textura,Calidad de piel | Zonas: Rostro,Cuello,Cuero cabelludo | Grado: Leve,Moderado | Sesiones: Según criterio médico | Intervalo: — | Recovery: Mínimo | Combinar: Microneedling,Peelings (según criterio) | NO combinar: — | Orden: Evaluación → toma → aplicación | CONTRAINDICACIONES: Embarazo,Infección activa,Alteración coagulación (según criterio médico)\n'
        '[Peelings Médicos] Precios: Melas Peel 3ses $200 / Vita C Peel 3ses $286 / Peeling Periocular 3ses $151 | Problemas: Manchas,Textura,Rejuvenecimiento facial | Zonas: Rostro | Grado: Leve,Moderado | Sesiones: Según protocolo | Intervalo: 2-6 semanas | Recovery: Variable | Combinar: Acthyderm,Hidrofacial (pre/post) | NO combinar: Exposición solar inmediata | Orden: Preparación → peeling → cuidado | CONTRAINDICACIONES: Embarazo,Infección activa,Piel lesionada,Isotretinoína reciente (según criterio médico)\n'
        '[Tratamientos Despigmentantes (Cosmelan / Melas Peel)] Precios: Cosmelan Kit $600 / Mantenimiento $300 / Melas Peel 3ses $200 | Problemas: Manchas,Melasma,Hiperpigmentación | Zonas: Rostro | Grado: Moderado,Severo | Sesiones: Según protocolo | Intervalo: — | Recovery: Variable | Combinar: Acthyderm,Hidrofacial (según plan) | NO combinar: Sol/broncedo reciente | Orden: Consulta → protocolo → control | CONTRAINDICACIONES: Embarazo,Infección activa,Piel lesionada\n'
        '[Capilar Plus] Precios: Capilar Plus 2ses $499 / Plasma Capilar 2ses $400 | Problemas: Caída de cabello,Calidad capilar | Zonas: Cuero cabelludo | Grado: Leve,Moderado | Sesiones: Según criterio médico | Intervalo: — | Recovery: Mínimo | Combinar: Plasma | NO combinar: — | Orden: Evaluación → plan → sesiones | CONTRAINDICACIONES: Infección activa,Alteración coagulación (según criterio médico)\n'
        '[Regenerador Facial] Precios: Regenerador Facial 3ses $613 / 1ses $313 | Problemas: Rejuvenecimiento facial,Poros/Textura,Arrugas/Líneas | Zonas: Rostro,Cuello | Grado: Moderado,Severo | Sesiones: 3-5 | Intervalo: 4-6 semanas | Recovery: 3-10 días | Combinar: Hidrofacial,Acthyderm | NO combinar: Peelings fuertes inmediato | Orden: Preparación → sesión → reparación | CONTRAINDICACIONES: Embarazo,Marcapasos,Infección activa,Alteración coagulación,Cáncer activo\n'
        '[Hidrofacial] Precios: Hidrofacial $90/ses | Problemas: Hidratación/Opacidad,Poros/Textura,Limpieza profunda | Zonas: Rostro | Grado: Leve,Moderado | Sesiones: 1-6 | Intervalo: 2-4 semanas | Recovery: Sin downtime | Combinar: Acthyderm,Peelings suaves | NO combinar: Peelings fuertes mismo día | Orden: Limpieza → activos | CONTRAINDICACIONES: Infección activa,Heridas abiertas (zona)\n'
        '[Microdermoabrasión] Precios: Microdermoabrasion $45/ses | Problemas: Poros/Textura,Opacidad,Manchas (superficiales) | Zonas: Rostro | Grado: Leve,Moderado | Sesiones: 1-6 | Intervalo: 2-4 semanas | Recovery: Rojeces leves | Combinar: Acthyderm,Hidrofacial | NO combinar: — | Orden: Exfoliación → activos | CONTRAINDICACIONES: Infección activa,Piel lesionada,Heridas abiertas\n'
        '[Peelings Cosméticos] Precios: Melas Peel 3ses $200 / Vita C Peel 3ses $286 | Problemas: Manchas,Poros/Textura,Opacidad | Zonas: Rostro | Grado: Leve,Moderado | Sesiones: 1-6 | Intervalo: 2-4 semanas | Recovery: Leve descamación | Combinar: Acthyderm,Hidrofacial | NO combinar: Sol inmediato | Orden: Preparación → peeling → hidratación | CONTRAINDICACIONES: Infección activa,Piel lesionada\n'
        '[Tratamientos con Péptidos] Precios: Peptidos Rejuvenecedores 3ses $544 / Parpados 3ses $187 | Problemas: Hidratación/Opacidad,Rejuvenecimiento facial | Zonas: Rostro | Grado: Leve,Moderado | Sesiones: 1-6 | Intervalo: 2-4 semanas | Recovery: Sin downtime | Combinar: Acthyderm,Hidrofacial | NO combinar: — | Orden: Limpieza → aplicación → sellado | CONTRAINDICACIONES: Alergias a componentes (según producto),Infección activa\n'
        '[Beauty Light / Fotofacial] Precios: Beauty Light 2ses $300 / Foto Facial 3ses $367 / Gleaming Skin 6ses $616 | Problemas: Manchas,Rejuvenecimiento facial,Enrojecimiento (según caso) | Zonas: Rostro,Cuello | Grado: Leve,Moderado | Sesiones: 3-6 | Intervalo: 4 semanas | Recovery: Leve enrojecimiento | Combinar: Hidrofacial,Acthyderm | NO combinar: Bronceado reciente | Orden: Preparación → IPL → hidratación | CONTRAINDICACIONES: Embarazo,Fotosensibilizantes,Infección activa,Sol/broncedo reciente\n'
        '[Bright Eyes] Precios: Bright Eyes 6ses $241 | Problemas: Ojeras,Textura,Rejuvenecimiento facial | Zonas: Rostro | Grado: Leve,Moderado | Sesiones: 1-6 | Intervalo: 2-4 semanas | Recovery: Sin downtime | Combinar: Acthyderm,Hidrofacial | NO combinar: — | Orden: Limpieza → protocolo ojos | CONTRAINDICACIONES: Infección activa,Heridas abiertas\n'
        '[Hidratación Piel Sensible] Precios: Hidratacion Piel 3ses $236 | Problemas: Hidratación/Opacidad,Reactividad | Zonas: Rostro | Grado: Leve,Moderado | Sesiones: 1-6 | Intervalo: 2-4 semanas | Recovery: Sin downtime | Combinar: Acthyderm | NO combinar: — | Orden: Limpieza suave → activos calmantes | CONTRAINDICACIONES: Alergias a componentes (según producto),Infección activa\n'
        '[Facial Personalizado] Precios: Blanqueamiento Facial 6ses $266 | Problemas: Hidratación/Opacidad,Poros/Textura,Rejuvenecimiento facial | Zonas: Rostro | Grado: Leve,Moderado | Sesiones: 1-6 | Intervalo: 2-4 semanas | Recovery: Sin downtime | Combinar: Acthyderm,Hidrofacial | NO combinar: — | Orden: Diagnóstico → protocolo | CONTRAINDICACIONES: Infección activa,Heridas abiertas\n'
        '[Lipoláser] Precios: Lipolaser 10ses $558 | Problemas: Grasa localizada,Contorno corporal | Zonas: Abdomen,Flancos,Piernas,Brazos | Grado: Leve,Moderado | Sesiones: 6-10 | Intervalo: 1 semana | Recovery: Sin downtime | Combinar: Presoterapia | NO combinar: — | Orden: Equipo → drenaje | CONTRAINDICACIONES: Embarazo,Marcapasos,Cáncer activo,Infección activa,Insuficiencia renal/hepática (según docs)\n'
        '[Electro Fit / Gimnasia Pasiva] Precios: Electro Fit 12ses $408 | Problemas: Tonificación,Flacidez corporal,Contorno corporal | Zonas: Abdomen,Glúteos,Piernas,Brazos | Grado: Leve,Moderado | Sesiones: 8-12 | Intervalo: 1 semana | Recovery: Sin downtime | Combinar: Presoterapia | NO combinar: — | Orden: Gimnasia → drenaje | CONTRAINDICACIONES: Embarazo,Marcapasos,Epilepsia,Hipertensión no controlada\n'
        '[Post Parto] Precios: Post Parto 10ses $218 | Problemas: Post parto,Retención de líquidos/Edema,Flacidez corporal | Zonas: Abdomen,Piernas | Grado: Leve,Moderado | Sesiones: 6-10 | Intervalo: 1 semana | Recovery: Sin downtime | Combinar: Gimnasia pasiva,Exilis | NO combinar: — | Orden: Drenaje → tonificación | CONTRAINDICACIONES: Embarazo,Marcapasos,Trombosis/TVP,Infección activa (según criterio)\n'
        '[Sculpted Body / Tensor Corporal] Precios: Sculped Body 12ses $458 / Tensor Cuerpo RF 8ses $808 | Problemas: Grasa localizada,Flacidez corporal,Contorno corporal | Zonas: Abdomen,Flancos,Brazos,Piernas | Grado: Moderado,Severo | Sesiones: 4-8 | Intervalo: 1-2 semanas | Recovery: Sin downtime | Combinar: Presoterapia,X-Wave | NO combinar: — | Orden: Equipo principal → drenaje → complementarios | CONTRAINDICACIONES: Embarazo,Marcapasos,Cáncer activo,Infección activa\n'
        '[Acthyderm Corporal] Precios: Acthyderm Cuerpo 12ses $783 | Problemas: Textura,Hidratación/Opacidad,Manchas | Zonas: Brazos,Cuerpo | Grado: Leve,Moderado | Sesiones: 4-8 | Intervalo: 1-2 semanas | Recovery: Sin downtime | Combinar: Exfoliación corporal | NO combinar: — | Orden: Exfoliación → activos → electroporación | CONTRAINDICACIONES: Embarazo,Marcapasos,Epilepsia,Cáncer activo,Infección activa\n'
        '[Ultra Focus Cuerpo] Precios: Total Lift paquete $1300 | Problemas: Flacidez corporal,Contorno corporal | Zonas: Abdomen,Brazos,Piernas | Grado: Moderado,Severo | Sesiones: 1-2 | Intervalo: 6-12 meses | Recovery: 24-72h leve edema | Combinar: Exilis,Presoterapia | NO combinar: Calor profundo intenso mismo día | Orden: Primero HIFU → complementarios | CONTRAINDICACIONES: Embarazo,Marcapasos,Cáncer activo,Infección activa,Heridas abiertas\n'
        '[Masajes & Relajación] Precios: Consultar | Problemas: Estrés,Tensión muscular,Bienestar | Zonas: Cuerpo,Rostro | Grado: Leve,Moderado | Sesiones: Según preferencia | Intervalo: — | Recovery: Sin downtime | Combinar: — | NO combinar: — | Orden: Sesión manual | CONTRAINDICACIONES: Lesión aguda en zona,Infección activa\n'
        '[Exfoliación Corporal] Precios: Blanqueamiento Corporal 6ses $266 | Problemas: Textura,Opacidad | Zonas: Cuerpo | Grado: Leve,Moderado | Sesiones: 1-6 | Intervalo: 2-4 semanas | Recovery: Sin downtime | Combinar: Acthyderm Corporal | NO combinar: — | Orden: Exfoliación → hidratación | CONTRAINDICACIONES: Heridas abiertas,Infección activa\n'
        '[Parafinas] Precios: Consultar | Problemas: Hidratación/Opacidad,Resequedad | Zonas: Manos/Pies,Espalda | Grado: Leve,Moderado | Sesiones: 1-6 | Intervalo: 2-4 semanas | Recovery: Sin downtime | Combinar: Masajes & Relajación | NO combinar: — | Orden: Aplicación → descanso | CONTRAINDICACIONES: Heridas abiertas,Infección activa,Sensibilidad al calor\n'
        '[Piernas Cansadas] Precios: Consultar | Problemas: Retención de líquidos/Edema,Piernas cansadas | Zonas: Piernas | Grado: Leve,Moderado | Sesiones: 4-8 | Intervalo: 1 semana | Recovery: Sin downtime | Combinar: Masajes & Relajación | NO combinar: — | Orden: Drenaje → descanso | CONTRAINDICACIONES: Trombosis/TVP,Insuficiencia arterial,Cáncer activo,Infección activa\n'
        '[Fisio Tape] Precios: Consultar | Problemas: Dolor muscular,Soporte postural,Bienestar | Zonas: Cuerpo | Grado: Leve,Moderado | Sesiones: Según criterio | Intervalo: — | Recovery: Sin downtime | Combinar: Masajes & Relajación | NO combinar: — | Orden: Evaluación → aplicación | CONTRAINDICACIONES: Alergia adhesivos,Heridas abiertas,Infección activa\n'
        '\nREGLAS: Verificar CONTRAINDICACIONES vs perfil del paciente. '
        'Respetar campo Orden y Recovery al planificar bimestres. '
        'Usar Combinar para sinergias. Respetar NO_combinar. '
        'total bimestre = suma real inversiones. total_anual = suma bimestres.'
    )

    if modelo == 'gemini':
        _llamar = _llamar_gemini
        tok1, tok2, tok3 = 6000, 12000, 10000
        nombre_modelo = 'Gemini'
    elif modelo == 'groq':
        _llamar = _llamar_groq
        tok1, tok2, tok3 = 8000, 12000, 10000  # más tokens = más detalle
        nombre_modelo = 'Groq (Llama)'
    else:
        _llamar = _llamar_claude
        tok1, tok2, tok3 = 8000, 12000, 10000
        nombre_modelo = 'Claude'

    actualizar(f'Sección 1/3 — Portada, diagnóstico y rutina diaria... ({nombre_modelo})', 15)
    r1, err = _llamar(1, 3, SYS1, datos, max_tok=tok1)
    if err: return {'error': err}

    if modelo == 'claude':
        time.sleep(2)  # pausa entre secciones para evitar overloaded (529)
    elif modelo in ('gemini', 'groq'):
        time.sleep(5)  # pausa para evitar rate limit en APIs con cuota baja

    actualizar(f'Sección 2/3 — Nutrición, ejercicio y bienestar mental... ({nombre_modelo})', 45)
    r2, err = _llamar(2, 3, SYS2, datos, max_tok=tok2)
    if err: return {'error': err}

    if modelo == 'claude':
        time.sleep(2)
    elif modelo in ('gemini', 'groq'):
        time.sleep(5)

    actualizar(f'Sección 3/3 — Sueño, tratamientos y plan de compromiso... ({nombre_modelo})', 75)
    r3, err = _llamar(3, 3, SYS3, datos, max_tok=tok3)
    if err: return {'error': err}

    resultado = {}
    resultado.update(r1)
    resultado.update(r2)
    resultado.update(r3)

    t_elapsed = round(time.time() - t_total, 1)
    print(f"[Total] {t_elapsed}s | claves: {list(resultado.keys())}")
    return resultado

# ════════════════════════════════════════════════════════════
# RENDER PLAN — llenar plantilla HTML
# ════════════════════════════════════════════════════════════

PLANTILLA_PLAN_HTML = """<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<title>Plan Carvajal · {{NOMBRE}}</title>
<link href="https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,300;0,400;0,600;1,300;1,400&family=Inter:wght@300;400;500;600;700;800&display=swap" rel="stylesheet">
<style>
:root{--olive:#8fa832;--olive-light:rgba(143,168,50,0.10);--olive-border:rgba(143,168,50,0.22);--dark:#1c1c1c;--gold:#b8935a;--cream:#faf9f6;--green:#2d3a2e;--white:#fff;--gray:#6b7280;--border:rgba(143,168,50,0.18)}
*{margin:0;padding:0;box-sizing:border-box}
body{font-family:"Inter",sans-serif;background:#d8d8d8;color:var(--dark);font-size:10.5pt;line-height:1.6}
h1,h2,h3,h4{font-family:"Cormorant Garamond",serif;line-height:1.2}
strong{font-weight:600}

/* ── TOPBAR ── */
#topbar-cv{position:fixed;top:0;left:0;right:0;z-index:9999;background:var(--green);padding:8px 20px;display:flex;align-items:center;justify-content:space-between;box-shadow:0 2px 12px rgba(0,0,0,0.3)}
#topbar-cv .tb-info strong{color:#fff;font-size:12px;display:block}
#topbar-cv .tb-info span{color:rgba(255,255,255,0.45);font-size:10px}
.tb-btns{display:flex;gap:8px}
.tb-btn{font-family:"Inter",sans-serif;font-size:11px;font-weight:500;cursor:pointer;padding:6px 14px;border-radius:4px;border:none;letter-spacing:.03em;transition:all .18s}
.tb-outline{background:transparent;border:1px solid rgba(255,255,255,.35);color:rgba(255,255,255,.85)}
.tb-outline:hover{border-color:#fff;color:#fff}
.tb-green{background:var(--olive);color:#fff}
.tb-green:hover{background:#7a9428}
.tb-green:disabled{background:#9ca3af;cursor:not-allowed}
#toast-cv{position:fixed;bottom:20px;right:20px;background:var(--green);color:#fff;padding:10px 18px;border-radius:5px;font-size:11px;font-family:"Inter",sans-serif;opacity:0;transition:opacity .3s;z-index:9998;pointer-events:none}
#toast-cv.show{opacity:1}
body{padding-top:44px}
[contenteditable]{outline:none;border-radius:2px;transition:background .15s}
[contenteditable]:hover:not(:focus){background:rgba(143,168,50,0.08)}
[contenteditable]:focus{background:rgba(143,168,50,0.13);box-shadow:0 0 0 1px rgba(143,168,50,0.35)}

@page{size:A4;margin:0}
@media print{
  #topbar-cv,#toast-cv,.no-print{display:none!important}
  body{background:white;padding-top:0}
  .page{box-shadow:none!important;margin:0!important;page-break-after:always}
  .page:last-child{page-break-after:auto}
  .pilar-card,.rutina-row,.suppl-item,.result-item,.chk-list li,.rut-step,.bim-body,.mini-card{page-break-inside:avoid}
}

/* ── PÁGINAS ── */
.page{width:210mm;min-height:297mm;margin:0 auto 8mm;background:white;position:relative;overflow:hidden;box-shadow:0 4px 40px rgba(0,0,0,0.15);display:flex;flex-direction:column}

/* PORTADA */
.cover{width:100%;height:297mm;background:white;display:flex;flex-direction:column;position:relative;overflow:hidden}
.cover-stripe{height:4px;background:var(--olive);width:100%;flex-shrink:0}
.cover-accent{position:absolute;top:0;right:0;width:3px;height:100%;background:linear-gradient(to bottom,var(--olive),transparent)}
.cover-top{padding:22px 44px;display:flex;align-items:center;justify-content:space-between;position:relative;z-index:2;background:#f7f7f7}
.cover-logo{height:72px;width:auto}
.cover-date-block{text-align:right}
.cdb-label{font-size:6.5pt;font-weight:600;text-transform:uppercase;letter-spacing:2.5px;color:var(--olive);margin-bottom:3px}
.cdb-date{font-size:9pt;font-weight:500;color:var(--gray)}
.cover-rule{height:1px;background:var(--olive-border);margin:16px 44px 0;position:relative;z-index:2}
.cover-title-block{padding:26px 44px 0;position:relative;z-index:2}
.cover-eyebrow{font-size:7pt;font-weight:600;text-transform:uppercase;letter-spacing:4px;color:var(--olive);margin-bottom:14px;display:flex;align-items:center;gap:10px}
.cover-eyebrow::after{content:"";flex:1;height:1px;background:var(--olive-border)}
.cover-main-title{font-family:"Inter",sans-serif;line-height:0.95;letter-spacing:-2px;margin-bottom:0}
.cover-main-title .t-light{display:block;font-weight:300;color:var(--gray);font-size:18pt;letter-spacing:-0.5px;margin-bottom:4px}
.cover-main-title .t-bold{display:block;font-weight:800;color:var(--olive);font-size:42pt;letter-spacing:-2.5px;line-height:0.92}
.cover-main-title .t-dark{display:block;font-weight:800;color:var(--dark);font-size:42pt;letter-spacing:-2.5px;line-height:0.92}
.cover-subtitle{font-family:"Inter",sans-serif;font-size:8.5pt;font-weight:400;color:var(--gray);letter-spacing:3px;text-transform:uppercase;margin-top:14px;padding-top:14px;border-top:1px solid var(--olive-border);display:inline-block}
.cover-patient-block{margin:24px 44px 0;padding:18px 22px;background:var(--dark);border-radius:6px;display:flex;align-items:center;gap:18px;position:relative;z-index:2}
.cpb-bar{width:3px;height:48px;background:var(--olive);border-radius:2px;flex-shrink:0}
.cpb-name{font-family:"Cormorant Garamond",serif;font-size:21pt;font-weight:600;color:white;line-height:1.1;margin-bottom:2px}
.cpb-meta{font-size:8pt;color:rgba(255,255,255,0.4)}
.cpb-data{margin-left:auto;display:flex;border-left:1px solid rgba(255,255,255,0.08);padding-left:18px}
.cpb-data-item{text-align:center;padding:0 16px;border-right:1px solid rgba(255,255,255,0.08)}
.cpb-data-item:last-child{border-right:none}
.cpdi-label{font-size:6pt;font-weight:600;text-transform:uppercase;letter-spacing:1.5px;color:var(--olive);margin-bottom:3px}
.cpdi-val{font-size:9.5pt;font-weight:600;color:white}
.cover-letter{margin:20px 44px 0;padding:16px 20px;background:var(--olive-light);border-left:3px solid var(--olive);border-radius:0 6px 6px 0;position:relative;z-index:2}
.cl-text{font-family:"Cormorant Garamond",serif;font-size:10.5pt;font-style:italic;line-height:1.75;color:var(--dark)}
.cover-pilares{margin:18px 44px 0;display:flex;flex-direction:column;gap:8px;position:relative;z-index:2}
.cp-row{display:flex;gap:8px}
.cp-row .cp-item{flex:1}
.cp-wide{flex:3!important}
.cp-narrow{flex:2!important}
.cp-item{background:white;border:1px solid var(--olive-border);border-top:3px solid var(--olive);border-radius:0 0 6px 6px;padding:20px 14px;text-align:center;min-height:90px;display:flex;flex-direction:column;align-items:center;justify-content:center}
.cp-icon{font-size:18pt;margin-bottom:7px}
.cp-num{font-size:6.5pt;color:var(--olive);font-weight:600;text-transform:uppercase;letter-spacing:1px;margin-bottom:4px}
.cp-label{font-size:7.5pt;font-weight:500;color:var(--dark);line-height:1.35}
.cover-spacer{flex:1}
.cover-footer{background:var(--dark);padding:12px 44px;display:flex;align-items:center;justify-content:space-between;position:relative;z-index:2;flex-shrink:0}
.cf-brand{font-size:8pt;font-weight:600;color:var(--olive)}
.cf-sub{font-size:7pt;color:rgba(255,255,255,0.28);margin-top:1px}
.cf-right{font-size:7pt;color:rgba(255,255,255,0.28);text-align:right;line-height:1.7}

/* PÁGINAS INTERNAS */
.page-header{background:var(--dark);padding:10px 28px;display:flex;align-items:center;justify-content:space-between;flex-shrink:0}
.ph-left{display:flex;align-items:center;gap:9px}
.ph-dot{width:6px;height:6px;border-radius:50%;background:var(--olive)}
.ph-title{font-size:8.5pt;color:var(--olive);font-weight:600}
.ph-right{font-size:6.5pt;color:rgba(255,255,255,0.22);text-transform:uppercase;letter-spacing:1px}
.content{padding:20px 28px;flex:1}
.content-sm{padding:14px 28px;flex:1}
.sec-label{font-size:6.5pt;font-weight:600;text-transform:uppercase;letter-spacing:3px;color:var(--olive);margin-bottom:4px}
.sec-title{font-family:"Cormorant Garamond",serif;font-size:19pt;font-weight:600;color:var(--dark);margin-bottom:15px;line-height:1.15}
.sec-title-sm{font-family:"Cormorant Garamond",serif;font-size:13pt;font-weight:600;color:var(--dark);margin-bottom:10px;padding-bottom:5px;border-bottom:1px solid var(--olive-border)}
.nota-medica{background:#fffbee;border-left:3px solid var(--gold);border-radius:0 6px 6px 0;padding:10px 14px;margin-bottom:13px;font-size:7.5pt;color:#7a5a10;line-height:1.6}
.nota-medica strong{color:var(--gold);display:block;margin-bottom:2px;font-size:7pt;text-transform:uppercase;letter-spacing:1px}
.diag-table{width:100%;border-collapse:collapse;font-size:7.5pt}
.diag-table th{background:var(--dark);color:var(--olive);padding:7px 11px;text-align:left;font-size:6.5pt;text-transform:uppercase;letter-spacing:1px;font-weight:600}
.diag-table td{padding:7px 11px;border-bottom:1px solid rgba(143,168,50,0.09);vertical-align:top;line-height:1.5}
.diag-table tr:nth-child(even) td{background:rgba(143,168,50,0.025)}
.diag-left{width:105px;font-weight:600;font-size:7.5pt;color:var(--dark);background:rgba(45,58,46,0.04)!important}
.badge-w{display:inline-block;background:#fef3cd;color:#856404;font-size:6.5pt;font-weight:600;padding:1px 5px;border-radius:6px}
.badge-c{display:inline-block;background:#f8d7da;color:#721c24;font-size:6.5pt;font-weight:600;padding:1px 5px;border-radius:6px}
.diag-val-strong{font-weight:600;font-size:8pt;color:var(--dark);display:block;margin-bottom:1px}
.diag-val-sub{font-size:7pt;color:var(--gray);line-height:1.5}
.pilar-card{display:flex;gap:11px;align-items:flex-start;padding:10px 13px;background:white;border:1px solid var(--olive-border);border-left:3px solid var(--olive);border-radius:0 6px 6px 0;margin-bottom:7px;page-break-inside:avoid}
.pilar-icon{font-size:14pt;flex-shrink:0;margin-top:2px}
.pilar-title{font-weight:600;font-size:8.5pt;color:var(--dark);margin-bottom:2px}
.pilar-desc{font-size:7.5pt;color:var(--gray);line-height:1.5}
.rutina-row{display:flex;align-items:flex-start;gap:11px;padding:7px 0;border-bottom:1px solid rgba(143,168,50,0.07);page-break-inside:avoid}
.rutina-hora{font-weight:700;color:var(--olive);min-width:40px;font-size:8.5pt;flex-shrink:0}
.rutina-text{flex:1;line-height:1.5;color:var(--dark);font-size:8pt}
.rtag{font-size:6pt;font-weight:600;padding:1px 7px;border-radius:3px;white-space:nowrap;flex-shrink:0;margin-top:2px;display:inline-block}
.rtag-n{background:#e8f5e9;color:#2e7d32}
.rtag-a{background:#e3f2fd;color:#1565c0}
.rtag-s{background:#f3e5f5;color:#4a148c}
.rtag-e{background:#fff8e1;color:#f57f17}
.rtag-h{background:#e8f5e9;color:#1b5e20}
.rtag-m{background:#f3e8ff;color:#4a148c}
.perm-evit{display:grid;grid-template-columns:1fr 1fr;gap:11px;margin-bottom:12px}
.pe-box{background:white;border:1px solid var(--border);border-radius:6px;padding:11px 13px}
.pe-title{font-size:7pt;font-weight:600;text-transform:uppercase;letter-spacing:1.5px;margin-bottom:7px;padding-bottom:5px;border-bottom:1px solid var(--border)}
.pe-title.p{color:#2e7d32}.pe-title.e{color:#c62828}
.pe-item{font-size:7.5pt;color:var(--dark);padding:2px 0;display:flex;align-items:flex-start;gap:6px;line-height:1.4}
.pe-dot{width:5px;height:5px;border-radius:50%;flex-shrink:0;margin-top:4px}
.pe-dot.p{background:#2e7d32}.pe-dot.e{background:#c62828}
.menu-table{width:100%;border-collapse:collapse;font-size:7pt}
.menu-table th{background:var(--olive);color:white;padding:5px 7px;text-align:left;font-size:6pt;text-transform:uppercase;letter-spacing:0.8px;font-weight:600}
.menu-table td{padding:5px 7px;border-bottom:1px solid rgba(143,168,50,0.09);vertical-align:top;line-height:1.4}
.menu-table .dia{font-weight:700;color:var(--olive);background:rgba(143,168,50,0.05)!important}
.suppl-item{display:flex;align-items:flex-start;gap:9px;padding:6px 0;border-bottom:1px solid var(--border);page-break-inside:avoid}
.suppl-item:last-child{border-bottom:none}
.suppl-bullet{width:18px;height:18px;border-radius:50%;background:var(--olive-light);border:1.5px solid var(--olive);display:flex;align-items:center;justify-content:center;font-size:7pt;font-weight:700;color:var(--olive);flex-shrink:0;margin-top:1px}
.suppl-name{font-weight:600;color:var(--dark);font-size:8pt}
.suppl-desc{font-size:7pt;color:var(--gray)}
.bim-header{background:var(--dark);color:var(--olive);padding:8px 13px;font-size:9pt;font-weight:600;border-radius:5px 5px 0 0;display:flex;align-items:center;justify-content:space-between}
.bim-body{border-radius:0 0 5px 5px;overflow:hidden;margin-bottom:9px;border:1px solid var(--border);border-top:none;page-break-inside:avoid}
.bim-table{width:100%;border-collapse:collapse;font-size:7pt}
.bim-table th{background:var(--olive-light);color:var(--olive);padding:5px 9px;text-align:left;font-size:6pt;text-transform:uppercase;letter-spacing:0.8px;font-weight:700}
.bim-table td{padding:6px 9px;border-bottom:1px solid var(--border);vertical-align:top;line-height:1.4}
.bim-table tr:last-child td{border-bottom:none}
.bim-total{background:var(--olive-light);padding:7px 13px;text-align:right;font-weight:700;color:var(--olive);font-size:8pt;border-top:1px solid var(--olive-border)}
.total-anual-box{background:var(--dark);border-radius:6px;padding:20px;text-align:center;margin:14px 0}
.ta-label{font-size:6.5pt;letter-spacing:3px;text-transform:uppercase;color:rgba(143,168,50,0.7);margin-bottom:5px}
.ta-amount{font-family:"Cormorant Garamond",serif;font-size:34pt;font-weight:600;color:var(--olive)}
.ta-sub{font-size:7pt;color:rgba(255,255,255,0.25);margin-top:3px}
.rut-grid{display:grid;grid-template-columns:1fr 1fr;gap:11px;margin-bottom:11px}
.rut-box{background:white;border:1px solid var(--border);border-radius:6px;overflow:hidden}
.rut-header{padding:8px 13px;font-size:8.5pt;font-weight:600;display:flex;align-items:center;gap:6px}
.rut-header.am{background:#fff8e1;color:#e65100;border-bottom:1px solid rgba(245,127,23,0.12)}
.rut-header.pm{background:#e8eaf6;color:#283593;border-bottom:1px solid rgba(57,73,171,0.12)}
.rut-step{display:flex;align-items:flex-start;gap:8px;padding:6px 11px;border-bottom:1px solid rgba(0,0,0,0.04);font-size:7.5pt;page-break-inside:avoid}
.rut-step:last-child{border-bottom:none}
.rut-num{width:17px;height:17px;border-radius:50%;background:var(--olive-light);color:var(--olive);display:flex;align-items:center;justify-content:center;font-size:6.5pt;font-weight:700;flex-shrink:0;margin-top:1px}
.rut-prod{font-weight:600;color:var(--dark);display:block;margin-bottom:1px}
.rut-desc{font-size:7pt;color:var(--gray)}
.cards-2{display:grid;grid-template-columns:1fr 1fr;gap:10px;margin-bottom:11px}
.cards-3{display:grid;grid-template-columns:repeat(3,1fr);gap:9px;margin-bottom:11px}
.mini-card{background:white;border:1px solid var(--border);border-radius:6px;padding:11px;page-break-inside:avoid}
.mc-icon{font-size:14pt;margin-bottom:5px}
.mc-title{font-size:8pt;font-weight:600;color:var(--dark);margin-bottom:3px}
.mc-text{font-size:7.5pt;color:var(--gray);line-height:1.5}
.objetivo-box{background:var(--green);border-radius:6px;padding:13px 17px;margin-bottom:12px}
.obj-label{font-size:6.5pt;font-weight:600;text-transform:uppercase;letter-spacing:2px;color:var(--olive);margin-bottom:4px}
.obj-text{font-size:8pt;color:rgba(255,255,255,0.85);line-height:1.65}
.quote-box{background:var(--olive-light);border-left:3px solid var(--olive);border-radius:0 6px 6px 0;padding:12px 16px;margin:11px 0}
.quote-text{font-family:"Cormorant Garamond",serif;font-style:italic;font-size:10pt;color:var(--olive);line-height:1.6}
.comp-box{background:var(--green);border-radius:6px;padding:17px 21px;margin-bottom:13px}
.comp-text{font-size:8pt;color:rgba(255,255,255,0.85);line-height:1.75}
.result-grid{display:grid;grid-template-columns:1fr 1fr;gap:7px;margin-bottom:13px}
.result-item{display:flex;align-items:flex-start;gap:8px;padding:8px 11px;background:white;border:1px solid var(--border);border-radius:5px;font-size:7.5pt;line-height:1.4;page-break-inside:avoid}
.result-check{width:17px;height:17px;background:var(--olive);border-radius:50%;display:flex;align-items:center;justify-content:center;flex-shrink:0;font-size:7.5pt;color:white;font-weight:700}
.chk-list{list-style:none}
.chk-list li{display:flex;align-items:flex-start;gap:9px;padding:8px 0;border-bottom:1px solid rgba(143,168,50,0.08);font-size:7.5pt;line-height:1.5;page-break-inside:avoid}
.chk-list li:last-child{border-bottom:none}
.chk-ico{width:17px;height:17px;border:1.5px solid var(--olive);border-radius:3px;display:flex;align-items:center;justify-content:center;flex-shrink:0;margin-top:1px}
.chk-ico svg{width:9px;height:9px;stroke:var(--olive);fill:none;stroke-width:2.5}
.divider-olive{height:2px;background:var(--olive);margin:12px 0;border-radius:1px;width:36px}
.page-footer{background:var(--dark);padding:7px 28px;display:flex;align-items:center;justify-content:space-between;font-size:6.5pt;color:rgba(255,255,255,0.25);flex-shrink:0}
.pf-brand{color:var(--olive);font-weight:600}
</style>
</head>
<body>

<!-- ══ TOPBAR (no se imprime) ══ -->
<div id="topbar-cv" class="no-print">
  <div class="tb-info">
    <strong>{{NOMBRE}}</strong>
    <span>Plan editable · {{FECHA}}</span>
  </div>
  <div class="tb-btns">
    <button class="tb-btn tb-outline" onclick="window.print()">🖨 Imprimir / PDF</button>
    <button class="tb-btn tb-green" id="btn-guardar" onclick="guardarCambios()">💾 Guardar cambios</button>
  </div>
</div>
<div id="toast-cv"></div>

<!-- ══ PORTADA ══ -->
<div class="page">
<div class="cover">
  <div class="cover-stripe"></div>
  <div class="cover-accent"></div>
  <div class="cover-top">
    <img src="https://centrocarvajal.com/wp-content/uploads/2023/05/logo-carvajal-300x147.jpg" class="cover-logo" alt="Centro Carvajal">
    <div class="cover-date-block">
      <div class="cdb-label">Fecha del plan</div>
      <div class="cdb-date" contenteditable="true">{{FECHA}}</div>
    </div>
  </div>
  <div class="cover-rule"></div>
  <div class="cover-title-block">
    <div class="cover-eyebrow">Programa de Salud, Bienestar y Desarrollo de Imagen</div>
    <div class="cover-main-title">
      <span class="t-light">Método de</span>
      <span class="t-bold">Rejuvenecimiento</span>
      <span class="t-dark">Carvajal</span>
    </div>
    <div class="cover-subtitle">Plan Integral Personalizado</div>
  </div>
  <div class="cover-patient-block">
    <div class="cpb-bar"></div>
    <div>
      <div class="cpb-name" contenteditable="true">{{NOMBRE}}</div>
      <div class="cpb-meta" contenteditable="true">{{EDAD}} años · {{OCUPACION}}</div>
    </div>
    <div class="cpb-data">
      <div class="cpb-data-item"><div class="cpdi-label">IMC</div><div class="cpdi-val" contenteditable="true">{{IMC}}</div></div>
      <div class="cpb-data-item"><div class="cpdi-label">Satisfacción</div><div class="cpdi-val" contenteditable="true">{{SATISFACCION}}/10</div></div>
      <div class="cpb-data-item"><div class="cpdi-label">Condición</div><div class="cpdi-val" contenteditable="true">{{CONDICION_CORTA}}</div></div>
    </div>
  </div>
  <div class="cover-letter">
    <div class="cl-text" contenteditable="true">{{INTRO}}</div>
  </div>
  <div class="cover-pilares">
    {{PILARES_PORTADA}}
  </div>
  <div class="cover-spacer"></div>
  <div class="cover-footer">
    <div>
      <div class="cf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</div>
      <div class="cf-sub">centrocarvajal.com · Tel: 263-8134 · 209-4284 · @centrocarvajal</div>
    </div>
    <div class="cf-right">Revisado por el equipo médico</div>
  </div>
</div>
</div>

<!-- ══ PÁG 2: DIAGNÓSTICO ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Diagnóstico y Análisis Inicial</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content">
    <div class="sec-label">Sección 01</div>
    <div class="sec-title">Diagnóstico Integral</div>
    <div style="border-radius:6px;overflow:hidden;border:1px solid var(--border)">
      <table class="diag-table">
        <thead><tr><th style="width:105px">Área</th><th>Hallazgos Clave</th></tr></thead>
        <tbody>{{DIAGNOSTICO_FILAS}}</tbody>
      </table>
    </div>
    <div style="margin-top:14px">
      <div class="sec-label" style="margin-bottom:7px">Los 5 Pilares del Plan</div>
      {{PILARES_CARDS}}
    </div>
  </div>
  <div class="page-footer"><span><span class="pf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</span></span><span>Página 2</span></div>
</div>

<!-- ══ PÁG 3: RUTINA + P1 NUTRICIÓN ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Rutina Diaria · Pilar 1: Nutrición</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content">
    <div class="sec-label">Sección 02</div>
    <div class="sec-title">Rutina Diaria Ideal</div>
    {{RUTINA_FILAS}}
    <div style="margin-top:16px">
      <div class="divider-olive"></div>
      <div class="sec-label" style="margin-bottom:7px">Pilar 1 · Sección 03</div>
      <div class="sec-title">{{P1_TITULO}}</div>
      <div class="objetivo-box">
        <div class="obj-label">Objetivo Focal</div>
        <div class="obj-text" contenteditable="true">{{P1_OBJETIVO}}</div>
      </div>
      <div class="perm-evit">
        <div class="pe-box">
          <div class="pe-title p">✓ Alimentos Recomendados</div>
          {{P1_PERMITIDOS}}
        </div>
        <div class="pe-box">
          <div class="pe-title e">✕ Limitar o Evitar</div>
          {{P1_EVITAR}}
        </div>
      </div>
    </div>
  </div>
  <div class="page-footer"><span><span class="pf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</span></span><span>Página 3</span></div>
</div>

<!-- ══ PÁG 4: MENÚ + SUPL + P2 ACTIVIDAD ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Menú Semanal · Suplementación · Actividad Física</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content-sm">
    <div class="sec-label">Pilar 1 · Continuación</div>
    <div class="sec-title" style="font-size:15pt;margin-bottom:9px">Menú Semanal</div>
    <div style="border-radius:6px;overflow:hidden;border:1px solid var(--border)">
      <table class="menu-table">
        <thead><tr><th style="width:52px">Día</th><th>Desayuno</th><th>Almuerzo</th><th>Cena</th><th>Snack</th></tr></thead>
        <tbody>{{P1_MENU}}</tbody>
      </table>
    </div>
    <div style="margin-top:12px">
      <div class="sec-label" style="margin-bottom:6px">Suplementación Recomendada</div>
      <div style="background:white;border:1px solid var(--border);border-radius:6px;padding:11px">
        {{P1_SUPLEMENTACION}}
      </div>
    </div>
    <div style="margin-top:12px">
      <div class="divider-olive"></div>
      <div class="sec-label" style="margin-bottom:5px">Pilar 2 · Sección 04</div>
      <div class="sec-title" style="font-size:14pt;margin-bottom:9px">{{P2_TITULO}}</div>
      <div class="objetivo-box" style="padding:11px 15px;margin-bottom:9px">
        <div class="obj-label">Objetivo Focal</div>
        <div class="obj-text" style="font-size:7.5pt" contenteditable="true">{{P2_OBJETIVO}}</div>
      </div>
      <div class="cards-2">
        <div class="mini-card">
          <div class="mc-title">Plan Semanal</div>
          <div class="mc-text" contenteditable="true">{{P2_PLAN}}</div>
        </div>
        <div class="mini-card">
          <div class="mc-title">Adaptaciones Específicas</div>
          <div class="mc-text" contenteditable="true">{{P2_ADAPTACIONES}}</div>
        </div>
      </div>
    </div>
  </div>
  <div class="page-footer"><span><span class="pf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</span></span><span>Página 4</span></div>
</div>

<!-- ══ PÁG 5: P3 BIENESTAR + P4 SUEÑO ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Pilar 3: Bienestar Mental · Pilar 4: Sueño</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content">
    <div class="sec-label">Pilar 3 · Sección 05</div>
    <div class="sec-title">{{P3_TITULO}}</div>
    <div class="objetivo-box">
      <div class="obj-label">Objetivo Focal</div>
      <div class="obj-text" contenteditable="true">{{P3_OBJETIVO}}</div>
    </div>
    <div class="cards-3">{{P3_TECNICAS_CARDS}}</div>
    <div class="quote-box">
      <div class="quote-text" contenteditable="true">{{P3_FRASE}}</div>
    </div>
    <div style="margin-top:14px">
      <div class="divider-olive"></div>
      <div class="sec-label" style="margin-bottom:5px">Pilar 4 · Sección 06</div>
      <div class="sec-title">{{P4_TITULO}}</div>
      <div class="objetivo-box" style="margin-bottom:12px">
        <div class="obj-label">Objetivo</div>
        <div class="obj-text" contenteditable="true">{{P4_OBJETIVO}}</div>
      </div>
      <div class="cards-2">
        <div class="mini-card">
          <div class="mc-title">Protocolo Nocturno</div>
          <div class="mc-text" contenteditable="true">{{P4_PROTOCOLO_TEXT}}</div>
        </div>
        <div class="mini-card">
          <div class="mc-title">Reglas Clave</div>
          <div class="mc-text" contenteditable="true">{{P4_REGLAS_TEXT}}</div>
        </div>
      </div>
    </div>
  </div>
  <div class="page-footer"><span><span class="pf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</span></span><span>Página 5</span></div>
</div>

<!-- ══ PÁG 6: TRATAMIENTOS BIM 1-4 ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Pilar 5 · Tratamientos — Bimestres 1 al 4</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content-sm">
    <div class="sec-label">Pilar 5 · Sección 07</div>
    <div class="sec-title" style="font-size:15pt">{{P5_TITULO}}</div>
    <div class="objetivo-box" style="padding:11px 15px;margin-bottom:11px">
      <div class="obj-label">Objetivo</div>
      <div class="obj-text" style="font-size:7.5pt" contenteditable="true">{{P5_OBJETIVO}}</div>
    </div>
    {{P5_BIMESTRES_A}}
  </div>
  <div class="page-footer"><span><span class="pf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</span></span><span>Página 6</span></div>
</div>

<!-- ══ PÁG 7: TRATAMIENTOS 5-6 + RUTINA FACIAL ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Tratamientos Bim. 5-6 · Total · Rutina Facial</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content-sm">
    {{P5_BIMESTRES_B}}
    <div class="total-anual-box">
      <div class="ta-label">Inversión Total del Plan</div>
      <div class="ta-amount">{{P5_TOTAL_ANUAL}}</div>
      <div class="ta-sub">12 meses · 6 bimestres · Plan integral personalizado · Centro Carvajal</div>
    </div>
    <div>
      <div class="sec-label" style="margin-bottom:5px">Rutina de Cuidado en Casa</div>
      <div class="rut-grid">
        <div class="rut-box">
          <div class="rut-header am">☀️ Rutina Mañana</div>
          {{P5_RUTINA_AM}}
        </div>
        <div class="rut-box">
          <div class="rut-header pm">🌙 Rutina Noche</div>
          {{P5_RUTINA_PM}}
        </div>
      </div>
      {{P5_NOTAS_CRITICAS}}
    </div>
  </div>
  <div class="page-footer"><span><span class="pf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</span></span><span>Página 7</span></div>
</div>

<!-- ══ PÁG 8: COMPROMISO ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Compromiso · Resultados · Próximos Pasos</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content">
    <div class="sec-label">Sección Final</div>
    <div class="sec-title">Compromiso y Seguimiento</div>
    <div class="comp-box">
      <div class="comp-text" contenteditable="true">{{COMP_PARRAFO}}</div>
    </div>
    <div class="sec-title-sm">Resultados Esperados a 12 Meses</div>
    <div class="result-grid">{{COMP_RESULTADOS}}</div>
    <div class="sec-title-sm" style="margin-top:12px">Próximos Pasos Inmediatos</div>
    <div style="background:white;border:1px solid var(--border);border-radius:6px;padding:11px">
      <ul class="chk-list">{{COMP_PASOS}}</ul>
    </div>
    <div class="quote-box" style="margin-top:12px">
      <div class="quote-text" contenteditable="true">{{COMP_FRASE}}</div>
    </div>
  </div>
  <div style="background:var(--dark);padding:12px 28px;text-align:center;flex-shrink:0">
    <div style="color:var(--olive);font-size:9pt;font-weight:600;margin-bottom:3px">Centro Carvajal · Líderes en Medicina Estética en Panamá</div>
    <div style="font-size:7pt;color:rgba(255,255,255,0.28);line-height:1.7">centrocarvajal.com · Tel: 263-8134 &amp; 209-4284 · @centrocarvajal · Panamá<br>Revisado y validado por el equipo médico de Centro Carvajal.</div>
  </div>
</div>

<script>
const JOB_ID = "{{JOB_ID}}";

async function guardarCambios() {
  const btn = document.getElementById("btn-guardar");
  btn.disabled = true; btn.textContent = "⏳ Guardando...";
  try {
    const resp = await fetch("/guardar/" + JOB_ID, {
      method: "POST",
      headers: {"Content-Type": "text/html; charset=utf-8"},
      body: document.documentElement.outerHTML
    });
    const data = await resp.json();
    if (data.ok) { showToast("✓ Guardado en Cloudinary"); btn.textContent = "✓ Guardado"; }
    else { showToast("Error: " + (data.error || "?"), 4000); btn.textContent = "💾 Guardar cambios"; }
  } catch(e) { showToast("Error de conexión", 4000); btn.textContent = "💾 Guardar cambios"; }
  finally { setTimeout(() => { btn.disabled = false; btn.textContent = "💾 Guardar cambios"; }, 2800); }
}

function showToast(msg, dur=2500) {
  const t = document.getElementById("toast-cv");
  t.textContent = msg; t.classList.add("show");
  setTimeout(() => t.classList.remove("show"), dur);
}
</script>
</body>
</html>"""

def render_plan(j, d, job_id=''):
    tpl = PLANTILLA_PLAN_HTML

    def esc(s): return str(s).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')

    nombre = d.get('nombre', '')

    # ── Diagnóstico ──
    badge_map = {'warning':'<span class="badge-w">⚠</span> ','critical':'<span class="badge-c">✕ Crítico</span> ','normal':''}
    diag_html = ''.join(
        f'<tr><td class="diag-left">{badge_map.get(f.get("alerta","normal"),"")}{esc(f.get("area",""))}</td><td><span class="diag-val-strong" contenteditable="true">{esc(f.get("estado",""))}</span><span class="diag-val-sub" contenteditable="true">{esc(f.get("hallazgos",""))}</span></td></tr>'
        for f in j.get('diagnostico', {}).get('filas', [])
    )

    # ── Portada: pilares masonry ──
    pr = j.get('portada', {}).get('pilares_resumen', [{}]*5)
    def get_p(i): return pr[i] if i < len(pr) else {}
    p = [get_p(i) for i in range(5)]
    portada_pilares = (
        f'<div class="cp-row">'
        f'<div class="cp-item cp-wide"><div class="cp-icon">{p[0].get("emoji","🥗")}</div><div class="cp-num">Pilar 1</div><div class="cp-label" contenteditable="true">{esc(p[0].get("titulo",""))}</div></div>'
        f'<div class="cp-item cp-narrow"><div class="cp-icon">{p[1].get("emoji","🏃")}</div><div class="cp-num">Pilar 2</div><div class="cp-label" contenteditable="true">{esc(p[1].get("titulo",""))}</div></div>'
        f'</div>'
        f'<div class="cp-row">'
        f'<div class="cp-item"><div class="cp-icon">{p[2].get("emoji","🧠")}</div><div class="cp-num">Pilar 3</div><div class="cp-label" contenteditable="true">{esc(p[2].get("titulo",""))}</div></div>'
        f'<div class="cp-item"><div class="cp-icon">{p[3].get("emoji","😴")}</div><div class="cp-num">Pilar 4</div><div class="cp-label" contenteditable="true">{esc(p[3].get("titulo",""))}</div></div>'
        f'<div class="cp-item"><div class="cp-icon">{p[4].get("emoji","✨")}</div><div class="cp-num">Pilar 5</div><div class="cp-label" contenteditable="true">{esc(p[4].get("titulo",""))}</div></div>'
        f'</div>'
    )

    # ── Pág 2: pilares cards ──
    pilares_cards = ''.join(
        f'<div class="pilar-card"><div class="pilar-icon">{pi.get("emoji","")}</div><div><div class="pilar-title" contenteditable="true">Pilar {pi.get("num","")} · {esc(pi.get("titulo",""))}</div><div class="pilar-desc" contenteditable="true">{esc(pi.get("descripcion",""))}</div></div></div>'
        for pi in j.get('portada', {}).get('pilares_resumen', [])
    )

    # ── Rutina ──
    tag_css = {'Nutricion':'rtag-n','Nutrición':'rtag-n','Sueno':'rtag-s','Sueño':'rtag-s','Actividad':'rtag-a','Mental':'rtag-m','Estetico':'rtag-e','Estético':'rtag-e','Salud':'rtag-h','Imagen':'rtag-e','Trabajo':'rtag-n'}
    rutina_html = ''.join(
        f'<div class="rutina-row"><div class="rutina-hora">{esc(r["hora"])}</div><div class="rutina-text" contenteditable="true">{esc(r["actividad"])}</div><div class="rtag {tag_css.get(r["pilar"],"rtag-n")}">{esc(r["pilar"])}</div></div>'
        for r in j.get('rutina', {}).get('items', [])
    )

    # ── P1 Nutrición ──
    p1 = j.get('pilar1', {})
    p1_perm = ''.join(f'<div class="pe-item"><div class="pe-dot p"></div><span contenteditable="true">{esc(i)}</span></div>' for i in p1.get('permitidos',[]))
    p1_evit = ''.join(f'<div class="pe-item"><div class="pe-dot e"></div><span contenteditable="true">{esc(i)}</span></div>' for i in p1.get('evitar',[]))
    p1_menu = ''.join(
        f'<tr><td class="dia">{esc(m.get("dia",""))}</td><td contenteditable="true">{esc(m.get("desayuno",""))}</td><td contenteditable="true">{esc(m.get("almuerzo",""))}</td><td contenteditable="true">{esc(m.get("cena",""))}</td><td contenteditable="true">{esc(m.get("snack",""))}</td></tr>'
        for m in p1.get('menu',[])
    )
    p1_supl = ''.join(
        f'<div class="suppl-item"><div class="suppl-bullet">{i+1}</div><div><div class="suppl-name" contenteditable="true">{esc(s)}</div></div></div>'
        for i, s in enumerate(p1.get('suplementacion',[]))
    )

    # ── P2 Actividad ──
    p2 = j.get('pilar2', {})

    # ── P3 Bienestar ──
    p3 = j.get('pilar3', {})
    icons3 = ['🧘','📵','📓','🌿','🛁','👩‍⚕️']
    p3_cards = ''.join(
        f'<div class="mini-card"><div class="mc-icon">{icons3[i] if i < len(icons3) else "✦"}</div><div class="mc-title" contenteditable="true">{esc(t)}</div></div>'
        for i, t in enumerate(p3.get('tecnicas',[]))
    )

    # ── P4 Sueño ──
    p4 = j.get('pilar4', {})
    p4_proto_text = '<br>'.join(f'<strong>{i+1}.</strong> {esc(s)}' for i, s in enumerate(p4.get('protocolo',[])))
    p4_reglas_text = '<br>'.join(f'• {esc(r)}' for r in p4.get('reglas',[]))

    # ── P5 Tratamientos (split bimestres A=1-4, B=5-6) ──
    p5 = j.get('pilar5', {})
    bimestres = p5.get('bimestres',[])

    def render_bim(bim):
        rows = ''.join(
            f'<tr><td contenteditable="true"><strong>{esc(t.get("nombre",""))}</strong></td><td contenteditable="true">{esc(t.get("sesiones",""))}</td><td contenteditable="true"><strong>{esc(t.get("inversion",""))}</strong></td><td contenteditable="true">{esc(t.get("beneficio",""))}</td></tr>'
            for t in bim.get('tratamientos',[])
        )
        total = bim.get('total',0)
        return (
            f'<div class="bim-header">{esc(bim.get("periodo",""))} · {esc(bim.get("titulo",""))}<span style="font-size:7pt;color:rgba(143,168,50,0.6)">Bimestre {bim.get("bimestre","")}</span></div>'
            f'<div class="bim-body"><table class="bim-table"><thead><tr><th>Tratamiento</th><th style="width:75px">Sesiones</th><th style="width:65px">Inversión</th><th>Beneficio</th></tr></thead><tbody>{rows}</tbody></table>'
            f'<div class="bim-total">💰 Inversión: ${total:,}</div></div>'
        )

    # Enumerate bimestres to add bimestre number
    for idx, bim in enumerate(bimestres):
        bim['bimestre'] = idx + 1

    half = max(4, len(bimestres) // 2 + len(bimestres) % 2)  # first 4 on page 6
    p5_bim_a = ''.join(render_bim(b) for b in bimestres[:half])
    p5_bim_b = ''.join(render_bim(b) for b in bimestres[half:])

    p5_am = ''.join(
        f'<div class="rut-step"><div class="rut-num">{s.get("paso","")}</div><div><div class="rut-prod" contenteditable="true">{esc(s.get("producto",""))}</div><div class="rut-desc" contenteditable="true">{esc(s.get("descripcion",""))}</div></div></div>'
        for s in p5.get('rutina_am',[])
    )
    p5_pm = ''.join(
        f'<div class="rut-step"><div class="rut-num">{s.get("paso","")}</div><div><div class="rut-prod" contenteditable="true">{esc(s.get("producto",""))}</div><div class="rut-desc" contenteditable="true">{esc(s.get("descripcion",""))}</div></div></div>'
        for s in p5.get('rutina_pm',[])
    )
    notas = p5.get('notas_criticas',[])
    p5_notas = ''
    if notas:
        items = ''.join(f'<span contenteditable="true">{esc(n)}</span><br>' for n in notas)
        p5_notas = f'<div class="nota-medica" style="margin-top:9px"><strong>⚠ Notas Críticas</strong>{items}</div>'

    total_anual = p5.get('total_anual', 0)

    # ── Compromiso ──
    comp = j.get('compromiso', {})
    comp_res = ''.join(
        f'<div class="result-item"><div class="result-check">✓</div><span contenteditable="true">{esc(r["texto"])}</span></div>'
        for r in comp.get('resultados',[])
    )
    comp_pasos = ''.join(
        f'<li><div class="chk-ico"><svg viewBox="0 0 24 24"><polyline points="20 6 9 17 4 12"/></svg></div><span contenteditable="true">{esc(p_)}</span></li>'
        for p_ in comp.get('proximos_pasos',[])
    )
    # Closing quote from any pilar frase
    comp_frase = p5.get('frase_motivacional') or p3.get('frase_motivacional') or ''

    # Condición corta para portada
    cond = d.get('condicionSistemica','') or ''
    cond_corta = cond[:20] + '…' if len(cond) > 20 else cond

    replacements = {
        '{{JOB_ID}}': job_id,
        '{{NOMBRE}}': esc(nombre),
        '{{EDAD}}': esc(d.get('edad','')),
        '{{OCUPACION}}': esc(d.get('ocupacion','')),
        '{{FECHA}}': esc(d.get('fecha','')),
        '{{IMC}}': esc(d.get('imc','N/A')),
        '{{SATISFACCION}}': esc(str(d.get('satisfaccion','?'))),
        '{{CONDICION_CORTA}}': esc(cond_corta),
        '{{INTRO}}': esc(j.get('portada',{}).get('intro','')),
        '{{PILARES_PORTADA}}': portada_pilares,
        '{{DIAGNOSTICO_FILAS}}': diag_html,
        '{{PILARES_CARDS}}': pilares_cards,
        '{{RUTINA_FILAS}}': rutina_html,
        '{{P1_TITULO}}': esc(p1.get('titulo','Nutrición')),
        '{{P1_OBJETIVO}}': esc(p1.get('objetivo','')),
        '{{P1_PERMITIDOS}}': p1_perm,
        '{{P1_EVITAR}}': p1_evit,
        '{{P1_MENU}}': p1_menu,
        '{{P1_SUPLEMENTACION}}': p1_supl,
        '{{P2_TITULO}}': esc(p2.get('titulo','Actividad Física')),
        '{{P2_OBJETIVO}}': esc(p2.get('objetivo','')),
        '{{P2_PLAN}}': esc(p2.get('plan_semanal','')),
        '{{P2_ADAPTACIONES}}': esc(p2.get('adaptaciones','')),
        '{{P3_TITULO}}': esc(p3.get('titulo','Bienestar Mental')),
        '{{P3_OBJETIVO}}': esc(p3.get('objetivo','')),
        '{{P3_TECNICAS_CARDS}}': p3_cards,
        '{{P3_FRASE}}': esc(p3.get('frase_motivacional','')),
        '{{P4_TITULO}}': esc(p4.get('titulo','Optimización del Sueño')),
        '{{P4_OBJETIVO}}': esc(p4.get('objetivo','')),
        '{{P4_PROTOCOLO_TEXT}}': p4_proto_text,
        '{{P4_REGLAS_TEXT}}': p4_reglas_text,
        '{{P5_TITULO}}': esc(p5.get('titulo','Tratamientos')),
        '{{P5_OBJETIVO}}': esc(p5.get('objetivo','')),
        '{{P5_BIMESTRES_A}}': p5_bim_a,
        '{{P5_BIMESTRES_B}}': p5_bim_b,
        '{{P5_TOTAL_ANUAL}}': f'${total_anual:,}',
        '{{P5_RUTINA_AM}}': p5_am,
        '{{P5_RUTINA_PM}}': p5_pm,
        '{{P5_NOTAS_CRITICAS}}': p5_notas,
        '{{COMP_PARRAFO}}': esc(comp.get('parrafo','')),
        '{{COMP_RESULTADOS}}': comp_res,
        '{{COMP_PASOS}}': comp_pasos,
        '{{COMP_FRASE}}': esc(comp_frase),
    }
    for k, v in replacements.items():
        tpl = tpl.replace(k, v)
    return tpl


def generar_calendario():
    mes_names = ['Enero','Febrero','Marzo','Abril','Mayo','Junio','Julio','Agosto','Septiembre','Octubre','Noviembre','Diciembre']
    dias_sem  = ['Lu','Ma','Mi','Ju','Vi','Sa','Do']
    hoy = date.today()
    yr  = hoy.year
    mi  = hoy.month
    cal_html = ''
    row_html = ''
    for m in range(12):
        mes     = ((mi - 1 + m) % 12) + 1
        anio_m  = yr + (mi - 1 + m) // 12
        dias_en = calendar.monthrange(anio_m, mes)[1]
        primer  = date(anio_m, mes, 1).isoweekday()
        grid    = '<div class="cal-grid">'
        for d in dias_sem: grid += f'<div class="cal-dh">{d}</div>'
        for _ in range(1, primer): grid += '<div class="cal-d cal-empty"></div>'
        for dia in range(1, dias_en + 1):
            grid += f'<div class="cal-d"><div class="n">{dia}</div><div class="dots"><div class="dot nu"></div><div class="dot ac"></div><div class="dot me"></div><div class="dot su"></div></div></div>'
        grid += '</div>'
        row_html += f'<div class="cal-month"><div class="cal-mhdr">{mes_names[mes-1]} {anio_m}</div>{grid}</div>'
        if (m + 1) % 3 == 0:
            cal_html += f'<div class="cal-row">{row_html}</div>'
            row_html = ''
    if row_html: cal_html += f'<div class="cal-row">{row_html}</div>'
    return cal_html


# ════════════════════════════════════════════════════════════
# EMAILS con Resend
# ════════════════════════════════════════════════════════════

def enviar_resend(asunto, cuerpo, to, adjunto_path=None, adjunto_name=None, adjuntos_extra=None, cc=None):
    if not RESEND_KEY:
        print('RESEND_KEY no configurado')
        return
    payload = {
        'from': f'Centro Carvajal <{MAIL_FROM}>',
        'to': [to],
        'subject': asunto,
        'html': cuerpo,
    }
    if cc:
        payload['cc'] = cc if isinstance(cc, list) else [cc]
    attachments = []
    if adjunto_path and adjunto_name and os.path.exists(adjunto_path):
        import base64
        with open(adjunto_path, 'rb') as f:
            attachments.append({
                'filename': adjunto_name,
                'content': base64.b64encode(f.read()).decode(),
            })
    for fp in (adjuntos_extra or []):
        if os.path.exists(fp):
            import base64
            with open(fp, 'rb') as f:
                attachments.append({
                    'filename': os.path.basename(fp),
                    'content': base64.b64encode(f.read()).decode(),
                })
    if attachments:
        payload['attachments'] = attachments
    try:
        print(f'[resend] Enviando a to={to} cc={payload.get("cc")} from={payload.get("from")} asunto={payload.get("subject","")[:50]}')
        r = req.post('https://api.resend.com/emails',
            headers={'Authorization': f'Bearer {RESEND_KEY}', 'Content-Type': 'application/json'},
            json=payload, timeout=30)
        print(f'[resend] Respuesta {r.status_code}: {r.text[:200]}')
    except Exception as e:
        print(f'Resend error: {e}')


def email_formulario(d, faltantes, borrador_url=''):
    nombre = d.get('nombre', '')
    rows = ''.join(f'<tr><td style="color:#b8935a;font-weight:600;padding:8px 16px;font-size:12px;text-transform:uppercase;letter-spacing:1px;width:140px">{k}</td><td style="padding:8px 16px;font-size:13px">{v}</td></tr>'
        for k, v in [('Nombre', nombre), ('Edad', d.get('edad','')), ('Ocupacion', d.get('ocupacion','')),
                     ('Medicamentos', d.get('medicamentos','')), ('Cirugias', d.get('cirugias','')),
                     ('Prioridad', d.get('prioridad','')), ('Satisfaccion', d.get('satisfaccion','') + '/10'),
                     ('Expectativas', d.get('expectativas',''))])
    faltantes_html = ''
    if faltantes:
        items = ''.join(f'<li style="font-size:12px;color:#6a5a20;padding:2px 0">{f}</li>' for f in faltantes)
        faltantes_html = f'<div style="background:#fffbf0;border:1px solid #e8d89a;border-radius:4px;padding:14px 18px;margin:16px 24px"><div style="font-size:11px;font-weight:700;color:#8a7030;margin-bottom:8px;text-transform:uppercase;letter-spacing:1px">Campos sin datos</div><ul style="padding-left:16px;margin:0">{items}</ul></div>'
    borrador_btn = f'<div style="text-align:center;margin:20px 24px"><a href="{borrador_url}" style="background:#8fa832;color:#fff;padding:13px 28px;border-radius:4px;text-decoration:none;font-size:14px;font-weight:500;display:inline-block">✏️ Revisar y editar borrador del plan</a><p style="font-size:11px;color:#999;margin-top:8px">Una vez editado, genera el PDF final desde el borrador.</p></div>' if borrador_url else ''
    return f'<!DOCTYPE html><html><head><meta charset="UTF-8"></head><body style="background:#f0e8de;padding:20px;font-family:sans-serif"><div style="max-width:600px;margin:0 auto;background:#fff;border:1px solid #ddd"><div style="background:#1a1410;padding:20px 24px"><div style="color:#b8935a;font-size:11px;letter-spacing:3px;text-transform:uppercase">Centro Carvajal · Nuevo Plan IA</div><div style="color:#fff;font-size:18px;margin-top:4px">{nombre}</div></div>{borrador_btn}<table style="width:100%;border-collapse:collapse">{rows}</table>{faltantes_html}<div style="background:#1a1410;padding:12px 24px;text-align:center;font-size:10px;color:rgba(255,255,255,0.3)">Centro Carvajal · centrocarvajal.com</div></div></body></html>'


def email_plan(nombre, html_url, fecha):
    return f'<!DOCTYPE html><html><head><meta charset="UTF-8"></head><body style="background:#f0e8de;padding:20px;font-family:sans-serif"><div style="max-width:600px;margin:0 auto;background:#fff;border:1px solid #ddd"><div style="background:#1a1410;padding:20px 24px"><div style="color:#b8935a;font-size:11px;letter-spacing:3px;text-transform:uppercase">Centro Carvajal · Plan Generado</div><div style="color:#fff;font-size:18px;margin-top:4px">{nombre}</div><div style="color:rgba(255,255,255,0.4);font-size:11px;margin-top:2px">{fecha}</div></div><div style="padding:24px"><p style="font-size:13px;color:#3d2e20;margin-bottom:16px">El plan personalizado de <strong>{nombre}</strong> ha sido generado exitosamente.</p><div style="text-align:center;margin:20px 0"><a href="{html_url}" style="background:#b8935a;color:#fff;padding:14px 32px;border-radius:4px;text-decoration:none;font-size:14px;font-weight:500">Ver Plan Completo</a></div><p style="font-size:11px;color:#999;text-align:center">O copia este link: {html_url}</p></div><div style="background:#1a1410;padding:12px 24px;text-align:center;font-size:10px;color:rgba(255,255,255,0.3)">Centro Carvajal · centrocarvajal.com</div></div></body></html>'


# ════════════════════════════════════════════════════════════
# MÓDULO DULCE DETALLE — esteticas.enmerida.mx/dulce-detalle
# ════════════════════════════════════════════════════════════

DD_MAIL = os.environ.get('DD_MAIL', os.environ.get('DEMO_MAIL', 'isai.josue@gmail.com'))

DULCE_CATALOGO = """
CATÁLOGO DULCE DETALLE MÉRIDA 2026
(Costo de envío NO incluido en ningún producto)

=== DESAYUNOS ===
- Presentación Básica — $280: Croissant de jamón de pavo y queso manchego, jugo del Valle, galleta, pretzel, etiqueta de ocasión.
- Chapata Básico — $280: Chapata de jamón con queso manchego, orejitas de hojaldre, chiles en raja, jugo.
- Fit Básico — $300: Fruta picada o croissant jamón/queso, jugo, yogurt, galletas, pretzel, granola.
- Box Sabritas — $320: Papas Sabritas, jugo, croissant jamón/queso, galletas, pretzels, chiles jalapeños. Caja cartón con etiqueta vinil.
- Charola Fit Básico — $340: Fruta picada, yogurt Oikos, granola, bebida energizante sin azúcar, agua, café Starbucks, pretzels.
- Box Ligero Mini / Desayuno Ligero Mini — $380: Croissant jamón/queso, fruta picada, jugo, pretzels, galletas, agua.
- Box Sabritas Krispy Kreme — $420: Croissant jamón/queso, dona Krispy Kreme, Sabritas, jugo, KitKat, galleta.
- Ligero Mini con Pastel — $460: Croissant jamón/queso, fruta picada, café frío, pastel mini, pretzels, galletas.
- Charola Sabritas Feliz Cumpleaños — $460: Croissant jamón/queso, Sabritas, pastel individual, jugo, galletas, pretzels.
- Batman Mini con Pastel — $480: Diseño Batman, vaso de colección, croissant jamón/queso, jugo, Sabritas, galletas, pretzels, pastel mini.
- Snoopy Mini con Pastel — $480: Diseño Snoopy, vaso de colección, croissant jamón/queso, jugo, Sabritas, galletas, pretzels, pastel mini.
- Astromelia Mini con Girasol — $520: Croissant jamón/queso, café frío, muffin, galletas, pretzels — con arreglo floral de astromelia y girasol.
- Chapata Special — $520: Chapata jamón/queso, orejitas de hojaldre, Ferreros (8 pzas), pretzels en frasco, café Starbucks, jugo, cubiertos.
- Astromelia — $580: Chapata jamón/queso, yogurt con fruta, orejitas, pretzels, jugo, cubiertos, taza con arreglo de astromelias.
- Chapata Black — $580: Chapata jamón/queso, pretzels en frasco, orejitas, Ferreros (8 pzas), café Starbucks, jugo, café frío, cubiertos.
- Romántico Black — $680: Chapata jamón/queso, ensalada frutas con yogurt/granola, pretzels, orejitas, jugo, café frío, cubiertos, base cerámica con 5 rosas rojas, charola de madera con etiqueta vinil.
- Isabella — $720: Chapata jamón/queso, arreglo floral rosas rojas + follaje, pretzels en frasco, orejitas, Ferreros (8 pzas), café Starbucks, jugo, café frío, cubiertos.
- Ligero Mini Premium — $880: Charola doble piso, arreglo floral 7 rosas o 2 girasoles, jugo, café frío, fruta picada, croissant jamón/queso, galletas, pretzels, KitKat. Etiqueta doble vinil.

=== BOX Y REGALOS ===
- Box Snoopy — $380: Croissant jamón/queso, vaso de colección Snoopy, jugo, galletas, pretzels.
- Box Corazón Compartido — $520: Croissant dulce fresa/Nutella + croissant salado jamón/queso, fresas, uvas, mini Nutella, galletas, pretzels, orejitas, pistaches, almendras, Ferreros.
- Box Starbucks — $520: Croissant jamón/queso Starbucks, bebida gasificada, Ferreros (8 pzas), KitKat, Waffle Starbucks.
- Box Starbucks y Rosas — $780: Arreglo floral rosas + astromelias, Agua Perrier, Waffle Starbucks, croissant jamón/queso, vaso reutilizable.

=== CANASTAS ===
- Canasta Desayuno Ligero Mini — $580: Canasta mimbre, arreglo floral, fruta picada (manzana, fresa, uva, papaya), jugo, croissant jamón/queso, galletas, pretzels.
- Canasta Starbucks Mini — $680: Croissant jamón/queso, café americano, dona Krispy Kreme, bebida gasificada, vaso de colección, arreglo floral girasol o 3 rosas.
- Canasta Starbucks Grande — $920: Croissant jamón/queso, fruta con yogurt, Ferreros (8 pzas), bebida gasificada, vaso Starbucks, Waffle Starbucks, arreglo floral girasol o 3 rosas.

=== CHAROLAS (DESAYUNO) ===
- Charola Love Chocolates — $580: Ferreros, 2 Kinder Delice, vaso Snoopy, Pulparindo, KitKat, Crunch, Paleta Payaso — con arreglo floral de rosas.
- Charola Starbucks Mini — $520: Croissant jamón/queso, café americano Starbucks, vaso de colección, bebida gasificada, dona Krispy Kreme, galletas, pretzels.
- Charola 6 Donas Krispy Kreme — $620: 6 donas glaseadas Krispy Kreme, bebida saborizada, café americano, galletas, pretzels, croissant jamón/queso, Ferreros (8 pzas).
- Charola Love Desayuno — $680: Arreglo floral 3 rosas + follaje, croissant jamón/queso, Cocacola lata, vaso Snoopy, galletas, pretzels.
- Charola de Frutas Premium — $1,100: Al menos 8 frutas de temporada en charola doble piso edición limitada, arreglo floral 7 rosas o 2 girasoles, etiquetas vinil.

=== FLORES Y ARREGLOS FLORALES ===
- Flor en Base de Cartón — $160
- Girasol con Astromelias y Tulia — $280
- Paquete Liss Girasol — $480
- Cinco Rosas y Tulia — $520
- Globo con Rosas y Mariposas — $580
- Corazón Doble Girasol y Astromelias — $580
- Rosas y Café Starbucks — $580
- Corazón de Ferreros y Rosas — $680
- Letra con Rosas y Chocolates — $720
- Corazón Rosas, Fresas y Chocolates — $720
- Corazón con Rosas y KitKat — $780
- Ramo de 24 Rosas Rosadas — $780
- Ramo Buchón 100 Rosas Rojas — $3,200

=== QUESOS Y CARNES FRÍAS ===
- Quesos y Carnes Frías Mini — $380
- Charola Carnes Frías y Quesos — $520
- Corazón de Madera con Carnes Frías — $620
- Tabla Corazón con Vino — $720
- Charola Carnes Frías y Cerveza — $820
- Charola de Carnes y Quesos — $920
- Carnes Frías Premium — $1,200
- Carnes Frías Premium y Rosas — $1,300

=== CHELAS, BOTELLAS Y BEBIDAS ===
- Cervezas y Botanas Mini — $420
- Charola Chelas y Botanas en Frascos — $620
- Box Cervezas Artesanales y Botanas — $680
- Charola Cervezas Artesanales — $880
- Tequila Don Julio y Botanas — $1,400
- Buchanans con Botanas — $1,800
- Tapete Rosas, Chocolates y Buchanan's — $2,600
"""

DULCE_PROMPT_BASE = (
    'Eres una asesora de regalos de Dulce Detalle Mérida — un negocio de desayunos sorpresa, flores, charcutería y arreglos frutales. '
    'Tu forma de comunicarte es cálida, cercana y genuina, como una amiga que te ayuda a elegir el regalo perfecto. '
    'Hablas de "tú". Usas el nombre de la persona. '
    '\n\nTONO — LA REGLA MÁS IMPORTANTE: '
    'Responde EXACTAMENTE al nivel emocional que el cliente expresó en su contexto. '
    'Si escribió poco o fue neutral → respuesta práctica, enfocada en el producto. '
    'Si fue emotivo y detallado → acompañar ese tono con calidez genuina. '
    'NUNCA asumas sentimientos que el cliente no expresó. '
    '\n\nRECUERDA SIEMPRE: el cliente está eligiendo un regalo para OTRA PERSONA. '
    '\n\nEVITA ABSOLUTAMENTE: '
    '"no dudes en contactarnos", "estamos a tu disposición", exclamaciones excesivas, lenguaje corporativo. '
    '\n\nESTRUCTURA: '
    '1. Abre con el nombre y una frase que demuestre que leíste su caso específico. '
    '2. Lista los productos recomendados — nombre en negrita + por qué encaja. '
    '3. Cierra con una invitación simple a coordinar la entrega. '
    '\n\nFORMATO HTML — solo estas etiquetas: <p>, <strong>, <ul>, <li>. '
    'NO uses markdown, NO asteriscos. NO menciones precios. '
    f'\n\nCATÁLOGO DISPONIBLE:\n{DULCE_CATALOGO}'
)


@app.route('/dulce-detalle')
def dulce_detalle():
    with open(os.path.join(os.path.dirname(__file__), 'index-dulce-detalle.html'), encoding='utf-8') as f:
        return f.read()


@app.route('/dulce-detalle/demo')
def dulce_detalle_demo():
    with open(os.path.join(os.path.dirname(__file__), 'formulario-dulce-detalle.html'), encoding='utf-8') as f:
        return f.read()


@app.route('/dulce-detalle/recomendar', methods=['POST'])
def dulce_recomendar():
    """Proxy multi-modelo para el formulario Dulce Detalle."""
    try:
        data   = request.get_json(force=True)
        perfil = data.get('perfil', '')
        modelo = data.get('modelo', 'claude')
        if not perfil:
            return jsonify({'error': 'Sin datos de perfil'}), 400

        user_msg = f'Por favor genera una recomendación personalizada para este cliente:\n\n{perfil}'

        # ── Prompt diferenciado por modelo ───────────────────
        if modelo in ('groq', 'openrouter'):
            sys_prompt = DULCE_PROMPT_BASE + (
                ' IMPORTANTE: sé muy específico y detallado — no te quedes en generalidades. '
                'Integra orgánicamente los datos reales del perfil: para quién es, el motivo, y especialmente el contexto adicional que escribió el cliente. '
                'Cada producto recomendado merece 2-3 oraciones: qué incluye, por qué encaja con este perfil específico y qué emoción o experiencia generará. '
                'Evita frases hechas como "no dudes en contactarnos", "estamos a tu disposición", "sería un placer", "el mejor curso de acción". '
                'Cierra con una sola oración natural invitando a coordinar la entrega, sin preguntas retóricas. '
                'Máximo 450 palabras.'
            )
        elif modelo == 'gemini':
            sys_prompt = DULCE_PROMPT_BASE + (
                ' IMPORTANTE: sé concreto y personalizado — cada recomendación debe sentirse escrita para esta persona, no para cualquiera. '
                'Menciona detalles reales del perfil del cliente en el texto de forma natural, especialmente lo que escribió en el contexto adicional. '
                'No uses listas genéricas — cada producto debe tener al menos una frase que explique por qué encaja con este caso en particular. '
                'Máximo 350 palabras. Cada frase debe aportar valor concreto — nada de relleno.'
            )
        else:
            sys_prompt = DULCE_PROMPT_BASE + ' Máximo 300 palabras. Cada frase debe aportar valor concreto — nada de relleno.'

        if modelo == 'claude':
            resp = req.post(
                'https://api.anthropic.com/v1/messages',
                headers={
                    'x-api-key': CLAUDE_KEY,
                    'anthropic-version': '2023-06-01',
                    'content-type': 'application/json'
                },
                json={
                    'model': 'claude-haiku-4-5-20251001',
                    'max_tokens': 900,
                    'system': sys_prompt,
                    'messages': [{'role': 'user', 'content': user_msg}]
                },
                timeout=30
            )
            resp.raise_for_status()
            texto = resp.json()['content'][0]['text']

        elif modelo == 'gemini':
            resp = req.post(
                f'https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-lite:generateContent?key={GEMINI_KEY}',
                headers={'Content-Type': 'application/json'},
                json={
                    'system_instruction': {'parts': [{'text': sys_prompt}]},
                    'contents': [{'parts': [{'text': user_msg}]}],
                    'generationConfig': {'maxOutputTokens': 900, 'temperature': 0.7}
                },
                timeout=30
            )
            resp.raise_for_status()
            texto = resp.json()['candidates'][0]['content']['parts'][0]['text']

        elif modelo == 'groq':
            resp = req.post(
                'https://api.groq.com/openai/v1/chat/completions',
                headers={'Authorization': f'Bearer {GROQ_KEY}', 'Content-Type': 'application/json'},
                json={
                    'model': 'llama-3.3-70b-versatile',
                    'max_tokens': 900,
                    'messages': [
                        {'role': 'system', 'content': sys_prompt},
                        {'role': 'user', 'content': user_msg}
                    ]
                },
                timeout=30
            )
            resp.raise_for_status()
            texto = resp.json()['choices'][0]['message']['content']

        elif modelo == 'openrouter':
            resp = req.post(
                'https://openrouter.ai/api/v1/chat/completions',
                headers={'Authorization': f'Bearer {OPENROUTER_KEY}', 'Content-Type': 'application/json'},
                json={
                    'model': 'nousresearch/hermes-3-llama-3.1-405b:free',
                    'max_tokens': 900,
                    'messages': [
                        {'role': 'system', 'content': sys_prompt},
                        {'role': 'user', 'content': user_msg}
                    ]
                },
                timeout=30
            )
            resp.raise_for_status()
            texto = resp.json()['choices'][0]['message']['content']

        else:
            return jsonify({'error': f'Modelo desconocido: {modelo}'}), 400

        texto = htmllib.unescape(texto)
        texto = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', texto)
        texto = re.sub(r'\*(.+?)\*', r'<em>\1</em>', texto)

        return jsonify({'html': texto})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/dulce-detalle/cita', methods=['POST'])
def dulce_cita():
    """Recibe solicitud de pedido y envía correo de notificación."""
    try:
        import resend as resend_lib
        resend_lib.api_key = RESEND_KEY

        d = request.get_json(force=True)
        nombre      = d.get('nombre', '')
        whatsapp    = d.get('whatsapp', '')
        email       = d.get('email', '')
        fecha       = d.get('fecha_entrega', 'No especificada')
        direccion   = d.get('direccion', 'No especificada')
        notas       = d.get('notas', '')
        para        = d.get('para', '')
        motivo      = d.get('motivo', '')
        tipo_regalo = d.get('tipo_regalo', '')
        presupuesto = d.get('presupuesto', '')

        cuerpo_negocio = f"""
<div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;padding:24px">
  <div style="background:#fff0f3;border:2px solid #f9a8b8;border-radius:16px;padding:20px 24px;margin-bottom:20px">
    <h2 style="font-size:22px;color:#2d2020;margin:0 0 4px">🎁 Nueva solicitud de pedido</h2>
    <p style="color:#9a7070;font-size:13px;margin:0">Dulce Detalle · esteticas.enmerida.mx/dulce-detalle</p>
  </div>
  <table style="width:100%;border-collapse:collapse;font-size:14px">
    <tr><td style="padding:10px 0;border-bottom:1px solid #f9c6d0;color:#9a7070;width:140px">👤 Nombre</td><td style="padding:10px 0;border-bottom:1px solid #f9c6d0;font-weight:600">{nombre}</td></tr>
    <tr><td style="padding:10px 0;border-bottom:1px solid #f9c6d0;color:#9a7070">📱 WhatsApp</td><td style="padding:10px 0;border-bottom:1px solid #f9c6d0;font-weight:600">{whatsapp}</td></tr>
    <tr><td style="padding:10px 0;border-bottom:1px solid #f9c6d0;color:#9a7070">✉️ Correo</td><td style="padding:10px 0;border-bottom:1px solid #f9c6d0;font-weight:600">{email}</td></tr>
    <tr><td style="padding:10px 0;border-bottom:1px solid #f9c6d0;color:#9a7070">🎁 Para quién</td><td style="padding:10px 0;border-bottom:1px solid #f9c6d0">{para}</td></tr>
    <tr><td style="padding:10px 0;border-bottom:1px solid #f9c6d0;color:#9a7070">🎉 Motivo</td><td style="padding:10px 0;border-bottom:1px solid #f9c6d0">{motivo}</td></tr>
    <tr><td style="padding:10px 0;border-bottom:1px solid #f9c6d0;color:#9a7070">🛍️ Tipo regalo</td><td style="padding:10px 0;border-bottom:1px solid #f9c6d0">{tipo_regalo}</td></tr>
    <tr><td style="padding:10px 0;border-bottom:1px solid #f9c6d0;color:#9a7070">💛 Presupuesto</td><td style="padding:10px 0;border-bottom:1px solid #f9c6d0">{presupuesto}</td></tr>
    <tr><td style="padding:10px 0;border-bottom:1px solid #f9c6d0;color:#9a7070">📅 Entrega</td><td style="padding:10px 0;border-bottom:1px solid #f9c6d0">{fecha}</td></tr>
    <tr><td style="padding:10px 0;border-bottom:1px solid #f9c6d0;color:#9a7070">📍 Dirección</td><td style="padding:10px 0;border-bottom:1px solid #f9c6d0">{direccion}</td></tr>
    <tr><td style="padding:10px 0;color:#9a7070">📝 Notas</td><td style="padding:10px 0">{notas or '—'}</td></tr>
  </table>
</div>
"""

        cuerpo_cliente = f"""
<div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;padding:24px">
  <div style="text-align:center;margin-bottom:24px">
    <div style="font-size:36px;margin-bottom:8px">🌸</div>
    <h2 style="font-family:Georgia,serif;font-size:24px;color:#2d2020;margin:0 0 6px">¡Hola, {nombre}!</h2>
    <p style="color:#9a7070;font-size:14px;margin:0">Recibimos tu solicitud en Dulce Detalle Mérida</p>
  </div>
  <div style="background:#fff0f3;border:2px solid #f9a8b8;border-radius:16px;padding:18px 22px;margin-bottom:20px">
    <p style="font-size:14px;color:#2d2020;line-height:1.7;margin:0">
      Gracias por confiar en nosotros para este regalo especial.
      Nos pondremos en contacto contigo pronto al <strong>{whatsapp}</strong> para coordinar todos los detalles.
    </p>
  </div>
  <table style="width:100%;border-collapse:collapse;font-size:14px">
    <tr><td style="padding:8px 0;border-bottom:1px solid #f9c6d0;color:#9a7070;width:120px">🎁 Para quién</td><td style="padding:8px 0;border-bottom:1px solid #f9c6d0">{para}</td></tr>
    <tr><td style="padding:8px 0;border-bottom:1px solid #f9c6d0;color:#9a7070">🎉 Motivo</td><td style="padding:8px 0;border-bottom:1px solid #f9c6d0">{motivo}</td></tr>
    <tr><td style="padding:8px 0;border-bottom:1px solid #f9c6d0;color:#9a7070">📅 Entrega</td><td style="padding:8px 0;border-bottom:1px solid #f9c6d0">{fecha}</td></tr>
  </table>
  <div style="text-align:center;margin-top:28px;padding-top:20px;border-top:2px dashed #f9c6d0">
    <p style="font-size:13px;color:#9a7070;margin:0">Dulce Detalle · Corporativo de Regalos · Mérida, Yucatán</p>
  </div>
</div>
"""

        resend_lib.Emails.send({
            'from': MAIL_FROM,
            'to': [DD_MAIL],
            'subject': f'🎁 Nuevo pedido — {nombre} ({motivo})',
            'html': cuerpo_negocio
        })

        if email:
            resend_lib.Emails.send({
                'from': MAIL_FROM,
                'to': [email],
                'subject': '🌸 Recibimos tu solicitud — Dulce Detalle Mérida',
                'html': cuerpo_cliente
            })

        return jsonify({'ok': True})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True, port=5000)
